from tkinter import *
from tkinter import ttk,messagebox
import tkinter as tk
from PIL import Image, ImageTk
import mysql.connector
from mysql.connector import Error
import pandas as pd
from docx.api import Document
#import aspose.words as aw
from tkinter import filedialog as fd
from tkcalendar import DateEntry
from datetime import date,timedelta,datetime
import xlrd #Version 1.2.0
import Permission
from ConnectToDB import *          #connect to mysql DB
import DB_tables                   #create tables

#from Permission import *
##table code
# https://pythonguides.com/python-tkinter-table-tutorial/

root = Tk()
# root.geometry("300x300")


root.title("Sheri Orders System")#Setting->to Main(Sheri Orders system)

# defult font
root.option_add("*Font", "Helvetica")


#validateLogin();

######################Hospital page##########################################

hospitalFrame = Frame(root)
h = Scrollbar(hospitalFrame, orient='horizontal')
#hospitalFrame.pack(fill=X)


# feed label
feedLabel = Label(hospitalFrame, text = 'Hospitals Details', font=('Helvetica',26, 'bold'),fg='#034672')
PlaceLable_X=50
PlaceLable_Y=10

feedLabel.pack(side=LEFT)
feedLabel.place(x=PlaceLable_X,y=PlaceLable_Y)

# admin label
AdminLabelHospitalpage1 = Label(hospitalFrame, text=f"{Permission.ValidateTypeOfUser} connected:", font=('Helvetica', 13, 'bold'), fg='red')
AdminLabelHospitalpage1.pack();
AdminLabelHospitalpage1.place(x=350, y=20);

# admin connected label
NameOfAdminLabelHospitalpage2 = Label(hospitalFrame, text=Permission.user_verified, font=('Helvetica', 13, 'bold'), fg='red')
NameOfAdminLabelHospitalpage2.pack();
NameOfAdminLabelHospitalpage2.place(x=510, y=20);



# scrollbar
Cyclotron_scroll = Scrollbar(hospitalFrame ,orient="vertical",width=25)
# Cyclotron_scroll.pack(side=LEFT)
# Cyclotron_scroll.place(x=550, y= 160)

hospitals_list = ttk.Treeview(hospitalFrame, yscrollcommand=Cyclotron_scroll.set,height=12)

hospitals_list.pack(side=LEFT, padx=PlaceLable_X+50, pady=PlaceLable_Y+80)

# column define

hospitals_list['columns'] = ('Name', 'Fixed Activity Level (mci)', 'Transport Time (minutes)')

# column format
width_Version=110
width_Capacity=110
width_Efficiency=185
width_Description=110

hospitals_list.column("#0", width=0, stretch=NO)
hospitals_list.column("Name", anchor=CENTER, width=width_Version)
hospitals_list.column("Fixed Activity Level (mci)", anchor=CENTER, width=width_Capacity)
hospitals_list.column("Transport Time (minutes)", anchor=CENTER, width=width_Efficiency)

# Create Headings
hospitals_list.heading("#0", text="", anchor=CENTER)
hospitals_list.heading("Name", text="Name", anchor=CENTER)
hospitals_list.heading("Fixed Activity Level (mci)", text="Fixed Activity Level (mci)", anchor=CENTER)
hospitals_list.heading("Transport Time (minutes)", text="Transport Time (minutes)", anchor=CENTER)

# add data from db
cursor = db.cursor();
cursor.execute("SELECT * FROM hospital");
hospitals_in_db = cursor.fetchall();

#Insert data of Hospitals into My-SQl DB
#The INSERT IGNORE statement will cause MySQL to do nothing when the insertion throws an error. If there’s no error, then a new row will be added to the table.
#cursor.execute("INSERT IGNORE INTO hospital (idhospital,Name,Fixed_activity_level,Transport_time) VALUES (1,'Belinson',9.2,15.0),(2,'Ichilov',10.0,20.0),(3,'Assuta TA',10.9,30.0),(4,'Sheb',10.5,35.0),(5,'Ziv',11.0,25.0),(6,'Assuta Ashdod',13.1,60.0),(7,'Assaf Harofeh',10.6,65.0),(8,'Augusta Victoria',9.6,50.0),(9,'Hila Pharma',9.6,50.0),(10,'Hadassah',9.5,0.0);")

#Insert 2 material to the DB,material table
#cursor.execute("INSERT IGNORE INTO material (idmaterial,materialName) VALUES (1,'FDG'),(2,'FDOPA');")
#cleanup DB
db.commit();
# cursor.close();
# db.close();

iid=0
for hospital in hospitals_in_db:
    #print(hospital)
    hospitals_list.insert(parent='', index='end', iid=iid, text='',
                          values=(hospital[1], hospital[2], hospital[3]))
    iid +=1

hospitals_list.pack()



# def open_popup_hospital():
#     pass
#
# def delete_hospital():
#     pass

#Create a button in the main Window to open the popup
global imgEdit;
editIcon = Image.open("D:/PythonProjects/Cyclotron/editIcon.jpg")
resizedEditIcon = editIcon.resize((20,20), Image.ANTIALIAS)
imgEdit = ImageTk.PhotoImage(resizedEditIcon,master=hospitalFrame)
editButton=Button(hospitalFrame,image=imgEdit, borderwidth=0)
editButton.pack()
editButton.place(x=425, y=55)


# delete button (Icon) - List
deleteIcon = Image.open("D:\PythonProjects\Cyclotron\‏‏deleteIcon.png")
resizedDeleteIcon = deleteIcon.resize((20,20), Image.ANTIALIAS)
imgDelete = ImageTk.PhotoImage(resizedDeleteIcon,master=hospitalFrame)
deleteButton=Button(hospitalFrame, image=imgDelete, borderwidth=0)
deleteButton.pack()
deleteButton.place(x=470, y=55)

#####################end of hospital page###################################################

# def hospital_page():
#     """ this function is swap function for hospital frame/page"""
#     #settingButton.config(bg='#F0F0F0');
#     #ordersButton.config(bg='#F0F0F0');
#     #hospitalsButton.config(bg='gray');
#     hospitalFrame.pack(fill='both',expand=1);
#     #SettingsFrame.forget();
#     ordersFrame.forget();

##################Orders page############################################

ordersFrame = Frame(root);
ordersFrame.config(bg="#F0F3F4");#color of page-white-gray

#ordersFrame.pack(fill=X)

# ScrollbarOrderMainPage
feedLabel = Label(ordersFrame, text ='Orders', font=('Helvetica', 26, 'bold'), fg='#034672');
feedLabel.pack(side=LEFT);
feedLabel.place(x=50,y=10);

# # admin label
AdminLabelOrderslpage1 = Label(ordersFrame, text=f"{Permission.ValidateTypeOfUser} connected:", font=('Helvetica', 13, 'bold'), fg='red')
AdminLabelOrderslpage1.pack();
AdminLabelOrderslpage1.place(x=350, y=20);
#
# # admin connected label
NameOfAdminLabeLoRDERSpage2 = Label(ordersFrame, text=Permission.user_verified, font=('Helvetica', 13, 'bold'), fg='red')
NameOfAdminLabeLoRDERSpage2.pack();
NameOfAdminLabeLoRDERSpage2.place(x=510, y=20);

# # user label
# UserLabelHospitalpage1 = Label(hospitalFrame, text=f"{Permission.ValidateTypeOfUser} connected:", font=('Helvetica', 13, 'bold'), fg='red')
# UserLabelHospitalpage1.pack();
# UserLabelHospitalpage1.place(x=350, y=20);
#
# # user  connected label
# NameOfUserLabelHospitalpage2 = Label(hospitalFrame, text=Permission.user_verified, font=('Helvetica', 13, 'bold'), fg='red')
# NameOfUserLabelHospitalpage2.pack();
# NameOfUserLabelHospitalpage2.place(x=490, y=20);


#my_label=Label(root,text='');

#Empty page/table for new order
OrdersTree = ttk.Treeview(ordersFrame,height=20);

# Defining number of columns
OrdersTree["columns"]=("1","2","3","4");

OrdersTree.pack(side=LEFT, padx=100, pady=110)

#Order main page scrollbar-vertical
OrderMainPagescroll = Scrollbar(ordersFrame, orient="vertical", command=OrdersTree.yview);
# OrderMainPagescroll.pack(side=RIGHT,fill=Y);
OrderMainPagescroll.place(x=705, y=110, height=330)

OrdersTree.configure(yscrollcommand = OrderMainPagescroll.set);


#Formate Columns
OrdersTree.column("#0",width=0,minwidth=0);
OrdersTree.column("1");
OrdersTree.column("2",width=120);
OrdersTree.column("3",width=40);
OrdersTree.column("4");


#Define headers/titles in table
OrdersTree.heading("#0", text="Label");
OrdersTree.heading("1", text="Hospital");
OrdersTree.heading("2", text="Injection Date");
OrdersTree.heading("3", text="Doses");
OrdersTree.heading("4", text="Last Updated");

#clear/delete Order main page tree-all records
def clear_tree():
    OrdersTree.delete(*OrdersTree.get_children());


#clear/delete Edit page tree-all records
def ClearEdittree_2():
    #"""Initialization-clear/Delete all the records"""
    for rawselected in EditTree.get_children():
        EditTree.delete(rawselected);

ListofCurrnetHospitalOrderMainPage=[];

global counterOrderId;

#if counterOrderId==0:
counterOrderId=0;

def updateOrdersTreeMainPageOutputOnly():

    clear_tree();
    # # Absorb Orders list data from db
    cursor = db.cursor();

    #Show output to Order main page tree-id,date,sum of doses,and last updated
    #cursor.execute("SELECT hospitalID,Date,SUM(amount) FROM orders GROUP BY Date,hospitalID;");
    cursor.execute("SELECT hospital.Name,orders.Date,SUM(orders.amount),MAX(updated_at) FROM orders INNER JOIN hospital ON hospital.idhospital = orders.hospitalID  GROUP BY orders.Date,orders.hospitalID;");
    SumOFAmount1 = cursor.fetchall();
    print(SumOFAmount1);
    # for x in SumOFAmount1:
    #     print(x);
    #convert list of tuples into list
    # ListofSumOfAmountPerHospital1 = [item1 for t1 in ListofSumOfAmountPerHospital for item1 in t1];

    # #Remove any duplicates from a List:-method 1
    # mylist = list(dict.fromkeys(ListofSumOfAmountPerHospital1))
    # ListOfSumOfAmount = [item1 for t1 in mylist for item1 in t1];

    # remove duplicated from list-method 2
    # FIxedHospitalList = []
    # for i in ListofCurrnetHospitalOrderMainPage:
    #     if i not in FIxedHospitalList:
    #         FIxedHospitalList.append(i)
    #
    # # printing list after removal
    # print ("The list after removing duplicates : " + str(FIxedHospitalList))

    #         cursor.execute("INSERT INTO orders (idorders,Date,injection_time,amount,hospitalID,batchID,DecayCorrected) VALUES (%s,%s,%s,%s,%s,%s,%s);",ValuseTuple);
    #output orders main data from DB to the orders tree
    for record in SumOFAmount1:
        OrdersTree.insert(parent='', index='end', text='',values=(record[0],record[1],record[2],record[3]))#record[0]=Idhospital,record[1]=Injection time,record[2]=Amount of doses,record[3]=last updated
        ListofCurrnetHospitalOrderMainPage.append(record[0]);

    #print(ListofCurrnetHospitalOrderMainPage);
    root.wm_state('normal');#Open orders main page
    #OrdersTree.pack();

    db.commit();
    cursor.close();

updateOrdersTreeMainPageOutputOnly();

def SearchOutpout(data):
    """Function for updating List of mian tree order main-output for searching component """

    clear_tree();
    # # Absorb Orders list data from db
    cursor = db.cursor();

    #Show output to Order main page tree-id,date,sum of doses
    #cursor.execute("SELECT hospitalID,Date,SUM(amount) FROM orders GROUP BY Date,hospitalID;");
    cursor.execute(f'SELECT hospital.Name,orders.Date,SUM(orders.amount),MAX(updated_at) FROM orders INNER JOIN hospital ON hospital.idhospital = orders.hospitalID GROUP BY orders.Date,orders.hospitalID;');
    DataAchievedBtSearch = cursor.fetchall();
    for x in DataAchievedBtSearch:
        print(x);
    #convert list of tuples into list
    # y = [item1 for t1 in SumOFAmount1 for item1 in t1];
    # print(y);
    #print(DataAchievedBtSearch);
    countertemp=0;
    for record in DataAchievedBtSearch:
        print(record)
        OrdersTree.insert(parent='', index='end', text='',values=(data[countertemp],record[1],record[2],record[3]));   #record[0]=hospital name,record[1]=Injection time,record[2]=Amount of doses,record[3]=updated_at
        countertemp=countertemp+1;


    root.wm_state('normal');#Open orders main page
    OrdersTree.pack();

    db.commit();
    cursor.close()



#########################Orders main pages buttons###############################

#Create Refresh(DB) button
global reafreahImg;
reafreshIcon = Image.open("D:/PythonProjects/Cyclotron/Images/regreshButton.png");
resizedReafreshEditIcon = reafreshIcon.resize((23,23), Image.ANTIALIAS);
reafreahImg = ImageTk.PhotoImage(resizedReafreshEditIcon);
ReafrshButton=Button(ordersFrame, image=reafreahImg, borderwidth=0,command=updateOrdersTreeMainPageOutputOnly)
ReafrshButton.pack();
ReafrshButton.place(x=530, y=63);


#Create Search window
searchEntry = Entry(ordersFrame,font=("Halvetica",12));
searchEntry.insert(0, 'Search Hospital Name');
searchEntry.pack();
searchEntry.place(x=640, y=65);

#Create search icon
searchIcon = Image.open("D:\PythonProjects\Cyclotron\Images\SearchButton.png");
resizedSearchedEditIcon = searchIcon.resize((23,23), Image.ANTIALIAS);
SearchImg = ImageTk.PhotoImage(resizedSearchedEditIcon);
SearchLabelicon=Label(image=SearchImg);
SearchLabelicon.pack();
SearchLabelicon.place(x=610, y=135);

def SearchComponent(event):
    """Function for creating search component """

    #get the Orders data from DB list-refresh
    cursor = db.cursor();

    #Show output to Order main page tree-id,date,sum of doses
    #cursor.execute("SELECT hospitalID,Date,SUM(amount) FROM orders GROUP BY Date,hospitalID;");
    cursor.execute("SELECT hospital.Name,orders.Date,SUM(orders.amount),MAX(updated_at) FROM orders INNER JOIN hospital ON hospital.idhospital = orders.hospitalID GROUP BY orders.Date,orders.hospitalID;");
    SumOFAmount1 = cursor.fetchall();
    print(SumOFAmount1);
    ListofCurrnetHospitalOrderMainPage=[];
    for record in SumOFAmount1:
        OrdersTree.insert(parent='', index='end', text='',values=(record[0],record[1],record[2],record[3])) #record[0]=Idhospital,record[1]=Injection time,record[2]=Amount of doses,record[3]=updated_at
        ListofCurrnetHospitalOrderMainPage.append(record[0]);

    #print(ListofCurrnetHospitalOrderMainPage);
    root.wm_state('normal');#Open orders main page
    #OrdersTree.pack();

    db.commit();
    cursor.close();



    typed = searchEntry.get();
    if typed == '':#if nothing typed in the entry
        dataList = ListofCurrnetHospitalOrderMainPage;#return all of curent list(relevant orders)
    else:
        dataList=[];
        for item in ListofCurrnetHospitalOrderMainPage:
            if typed.lower() in item.lower():           #catch capital and lower case
                dataList.append(item);
                print("Hospital finded:",dataList);

    SearchOutpout(dataList);#update past list in the new searched list

searchEntry.bind("<KeyRelease>",SearchComponent)#catch any key pressed and released from keyboard- event

#Clear search button on-click in the input widget
# define an event handler
def handleEvent(event):
    searchEntry.delete(0, END);

# bind the click of the Entry to the handler
searchEntry.bind("<ButtonRelease>", handleEvent)


#######Filter by material-drop down##################################################
# MaterialClickDropDownMenu = StringVar();
# MaterialClickDropDownMenu.set("Material"); #default value
#
# Absorb materials list data from db
cursor = db.cursor();
cursor.execute("SELECT idmaterial,materialName FROM material");
Material_in_db = cursor.fetchall();
print(Material_in_db);
AllListOption=("All");
Material_in_db.append(AllListOption);
print(Material_in_db);
def updateOrdersTreeByMaterialFiltering(materialSelData):
    clear_tree();
    # # Absorb Orders list data from db
    cursor = db.cursor();
    if materialSelData=='A':
        #Show output to Order main page tree-id,date,sum of doses
        cursor.execute(f"SELECT hospital.Name,orders.Date,SUM(orders.amount),MAX(updated_at) FROM orders  INNER JOIN hospital ON hospital.idhospital = orders.hospitalID GROUP BY orders.Date,orders.hospitalID ;");
    else:
        #Show output to Order main page tree-id,date,sum of doses filtering by Material ID
        cursor.execute(f"SELECT hospital.Name,orders.Date,SUM(orders.amount),MAX(updated_at) FROM orders  INNER JOIN hospital ON hospital.idhospital = orders.hospitalID WHERE materialID={materialSelData}  GROUP BY orders.Date,orders.hospitalID ;");

    filteringRowsFromDB = cursor.fetchall();
    print(filteringRowsFromDB);
    for record in filteringRowsFromDB:
        OrdersTree.insert(parent='', index='end', text='',values=(record[0],record[1],record[2],record[3]))#record[0]=Idhospital,record[1]=Injection time,record[2]=Amount of doses,record[3]=last updated
        ListofCurrnetHospitalOrderMainPage.append(record[0]);

    #print(ListofCurrnetHospitalOrderMainPage);
    root.wm_state('normal');#Open orders main page
    OrdersTree.pack();

    db.commit();
    cursor.close();


# #function to get the value from the drop down menu
def MaterialsSelectedeFilteringFunc(event):
    #MaterialDropDownLabel=Label(ordersFrame,text=MaterialsDropDownFilteringMainPage.get()).pack()
    MaterialSelected=MaterialsDropDownFilteringMainPage.get();
    Mlist=list(MaterialSelected);
    print(Mlist[0]);
    updateOrdersTreeByMaterialFiltering(Mlist[0]);#Calling to output sql function


MaterialsDropDownFilteringMainPage = ttk.Combobox(ordersFrame,state="readonly",value=Material_in_db,width=9);
MaterialsDropDownFilteringMainPage.current(2);

MaterialsDropDownFilteringMainPage.bind("<<ComboboxSelected>>",MaterialsSelectedeFilteringFunc)
MaterialsDropDownFilteringMainPage.pack();
MaterialsDropDownFilteringMainPage.place(x=370, y=70);

############################################################################

# Absorb hosital list data from db
# cursor = db.cursor();
# cursor.execute("SELECT idhospital,Name FROM hospital");
# hospitals_in_db = cursor.fetchall();
# HospitalListForDeleteOrder = hospitals_in_db;
# print(HospitalListForDeleteOrder);


OrderselectedEvent = tk.StringVar();
def deleteOrderEvent(event):

    """Function for removing order from DB"""
    global DoseNum_Selected,DateSelected,InjectionTimeSelected,IDofHospitalSelected2;
    row = OrdersTree.focus();

    dataofchoosnenRowListEditTree=row;
    #print(dataofchoosnenRowListEditTree);
    DataOfRowSelectedDicEditTree=OrdersTree.item(dataofchoosnenRowListEditTree);
    DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
    print("Record/Order selected to delete: ",DataOfRowSelectedList);
    hospitalSelected=DataOfRowSelectedList[0];
    DateSelected=DataOfRowSelectedList[1];
    InjectionTimeSelected=DataOfRowSelectedList[2];

    #search hospital by name from hospital table db and get the ID as output
    cursor = db.cursor();
    cursor.execute(f'SELECT CAST(idhospital AS SIGNED) FROM hospital WHERE Name="{hospitalSelected}"');
    IDofHospitalSelected1 = cursor.fetchall();
    TempID=[i[0] for i in IDofHospitalSelected1];#find index number in a list of tuple
    IDofHospitalSelected2=int(TempID[0]);
    print(f'{IDofHospitalSelected2} : {hospitalSelected}');

def deleteOrderfunc():
    MsgBox = messagebox.askquestion ('Info message','Are you sure you want to cancel/delete the order?',icon = 'warning')
    if MsgBox == 'yes':
        rawSelectedToDelete=OrdersTree.selection();#selected item:I001,I002,I003....
        RecoredDeletedFlug=1;
        try:
            cursor = db.cursor(buffered=True);
            for rawselected in rawSelectedToDelete:
                UpdateSQlQuery=f"UPDATE  orders SET deleted='{RecoredDeletedFlug}' WHERE  hospitalID= '{IDofHospitalSelected2}' AND Date= '{DateSelected}';";
                DeleteQuery = f"DELETE FROM orders WHERE hospitalID= '{IDofHospitalSelected2}' AND Date= '{DateSelected}';";
                cursor.execute(UpdateSQlQuery);
                cursor.execute(DeleteQuery);
                OrdersTree.delete(rawselected);
                print("DB updated successfully-Record add to deleted column ");
                db.commit();
                cursor.close();
        except mysql.connector.errors.IntegrityError as e:
            messagebox.showinfo("info message","Watch out\nThere is workplan that relate to the order you trying to delete/cancel that still active\n,action cancel!");
            logging.error(traceback.format_exc())
        except Exception as e:
            logging.error(traceback.format_exc())
            print("Error-Order was not updated-please check MySQL");

    else:
        messagebox.showinfo('Return','Return to the Orders screen');



# Remove button (Icon) -Delete Order
global imgDelete2;
deleteIcon = Image.open("D:\PythonProjects\Cyclotron\‏‏deleteIcon.png")
resizedDeleteIcon = deleteIcon.resize((20,20), Image.ANTIALIAS)
imgDelete2 = ImageTk.PhotoImage(resizedDeleteIcon)
deleteButton=Button(ordersFrame, image=imgDelete2, borderwidth=0,command=deleteOrderfunc)
deleteButton.pack()
deleteButton.place(x=560, y=65)

OrdersTree.bind('<<TreeviewSelect>>', deleteOrderEvent);



###############################################Import File page##################################


def WriteToCsv(result):
    """Function for creating/exporting Excel file"""
    print("try exporting new excel file...");
    headers = ['OrderId', 'Date', 'Injection Time', 'Amount','idhospital','batchID','decayCorrected'];
    with open('orders.csv','a',newline="") as f:
        w = csv.writer(f,dialect='excel');
        messagebox.showinfo("message","Excel file was created");
        # write the headers
        w.writerow(headers);
        for record in result:
            w.writerow(record);


# Absorb Orders table data from db-for excel export
cursor = db.cursor();
cursor.execute("SELECT * FROM orders");
ordersTable_in_db = cursor.fetchall();

#Create Export to Excel buttton
global ExportToCSVImg;

ExportCSVIcon = Image.open("D:\PythonProjects\Cyclotron\Images\ExportExcel.png");
resizedExportCSVIcon = ExportCSVIcon.resize((23,23), Image.ANTIALIAS);
ExportToCSVImg = ImageTk.PhotoImage(resizedExportCSVIcon);
ExportToCSVImgicon=Button(ordersFrame, image=ExportToCSVImg, borderwidth=0,command=lambda : WriteToCsv(ordersTable_in_db))
ExportToCSVImgicon.pack();
ExportToCSVImgicon.place(x=585, y=63);

#Create edit icon
# global imgEdit;
# editIcon = Image.open("editIcon.jpg")
# resizedEditIcon = editIcon.resize((20,20), Image.ANTIALIAS)
# imgEdit = ImageTk.PhotoImage(resizedEditIcon)
# editButton=Button(ordersFrame, image=imgEdit, borderwidth=0)
# editButton.pack()
# editButton.place(x=800, y=65)
# edit_button = Button(ordersFrame, text= "Edit")
# edit_button.pack(side= LEFT)
# edit_button.place(x=450, y=50)

#edit field from DB
# query = "UPDATE hospital SET Name = %s ,Fixed_activity_level= %s, Transport_time=%s  WHERE idhospital = %s"
#
# pk = selected_rec[3]
#
# labels = (('Name', ''), ('Fixed activity level', '(mci/h)'),  ('Transport time', '(min)'))
# save_title = "Save Changes"
#
# editHospitalPopup.edit_popup(labels, selected_rec, save_title, query, pk, hospital_tabel)

########################Import File page##################################################################
def importFileFunc():
    AmountListFromDoc=[];
    InjectionTImeListFromdoc=[];
    #ListofVarImportFile=["","","","","",""];
    TempList=["","",""];
    def ImportFilefunction():
        TempNewLISt1=[];
        """This is function for importing Orders files"""
        filename = fd.askopenfilename(
            initialdir="D:\PythonProjects\Cyclotron",
            title="Open a file",
            filetype=(("Word files","*.docx"),("Word files","*.doc"),("Excel files","*.xlsx"),("All Files","*.*"),("PDF files","*.pdf")))
        #print(filename);
        if filename:
            if  "xlsx" in filename :                     #Excel file
                try:
                    filename=r"{}".format(filename)
                    df=pd.read_excel(filename)
                except ValueError:
                    messagebox.showinfo("Error message","File couldn't be open,try again");
                    print("Error");
                except FileNotFoundError:
                    messagebox.showinfo("Error message","File couldn't be open,try again");
                    print("Error");

                #clear_tree();


                #####################################################################
                #Get data from excel
                loc = (str(filename));
                wb = xlrd.open_workbook(loc);
                sheet = wb.sheet_by_index(0);
                sheet.cell_value(0, 0);
                #Get amount data/column from doc/excel-Get amount
                for i in range(1,sheet.nrows):
                    AmountListFromDoc.append(sheet.cell_value(i, 2));
                    print(f"Amount number{i}:{AmountListFromDoc[i-1]}");

                    #Get Injection time data/column from doc/excel-Get Injection time
                for i in range(1,sheet.nrows):
                    InjectionTImeListFromdoc.append(sheet.cell_value(i, 5));
                    print(f'InjectionTIme numner {i}: {InjectionTImeListFromdoc[i-1]}');
                ########################################################

            #word files

            if "docx" in filename or "doc" in filename:
                #convert word to excel

                if (("doc" in filename) and ("docx" not in filename)):#convert docx to doc
                    doc = aw.Document(filename)
                    filename="NewWordOutput1.docx";
                    doc.save(filename)


                document = Document(filename);
                tables = document.tables;
                df = pd.DataFrame();

                for table in document.tables:
                    for row in table.rows:
                        text = [cell.text for cell in row.cells];
                        df = df.append([text], ignore_index=True);

                #df.columns = ["Column1", "Column2","Column3","Column4","Column5","Column6","Column7","Column8"]
                df.to_excel("D:/PythonProjects/Cyclotron/OrderOutputTest.xlsx");
                #print(df);


                #clear_tree();
                #Get data from excel
                loc = ("D:/PythonProjects/Cyclotron/OrderOutputTest.xlsx");
                wb = xlrd.open_workbook(loc);
                sheet = wb.sheet_by_index(0);
                sheet.cell_value(0, 0);
                #Get amount data/column from doc/excel
                for i in range(1,sheet.nrows):
                    AmountListFromDoc.append(sheet.cell_value(i, 2));
                    print(AmountListFromDoc)

                    #Get Injection time data/column from doc/excel-Get Injection time
                for i in range(1,sheet.nrows):
                    InjectionTImeListFromdoc.append(sheet.cell_value(i, 5));
                    print(f'InjectionTIme numner {i}: {InjectionTImeListFromdoc[i-1]}');
        # ##############################################################

        # FileImportedLabel1=Label(ImportFilePage, text="FIle imported successfully",fg="red", font=('Helvetica 12'));
        # FileImportedLabel1.pack();
        # FileImportedLabel1.place(x=450,y=180);

        FileImportedLabel2=Label(ImportFilePage, text=filename, font=('Helvetica 12'),fg="red");
        FileImportedLabel2.pack();
        FileImportedLabel2.place(x=450,y=200);

        root.wm_state('iconic');#minimize orders main page



    ##########################################
    ImportFilePage = Toplevel(root);
    ImportFilePage.geometry("900x400");
    ImportFilePage.config(bg="#F0F3F4");#color of page-white-gray
    ImportFilePage.title("Import File");
    #NewOrdersecondaryPage = tk.Frame(root);

    ImpoerFilePageLabel=Label(ImportFilePage, text="Import File - Order", font=('Helvetica 17 bold'), fg='#034672');
    ImpoerFilePageLabel.pack();

    HospitalListLabel = Label(ImportFilePage, text="Hospital",bg='white');
    HospitalListLabel.pack();
    HospitalListLabel.place(x=20, y=70);

    #Create hospital drop-down
    # Absorb hosital list data from db
    cursor = db.cursor();
    cursor.execute("SELECT idhospital,Name FROM hospital");
    hospitals_in_db = cursor.fetchall();
    HospitalList2 = hospitals_in_db;


    def HospitalChooseImportFile(HospitalSelectedEvent):
        """Function for create Hospital Drop-Down menu -absorb data from DB"""
        ChoosenHospitalNewOrder=HospitalSelectedImportFile.get();
        print(ChoosenHospitalNewOrder);

        #loop for findinf Id number in the string
        HospitalIDFromChoosenHospital = "";
        for index in ChoosenHospitalNewOrder:
            if index.isdigit():
                HospitalIDFromChoosenHospital = HospitalIDFromChoosenHospital + index;


        TempList[0]=HospitalIDFromChoosenHospital;#HospitalID
        print(TempList[0]);



    HospitalSelectedImportFile = ttk.Combobox(ImportFilePage,state="readonly",value=HospitalList2,width=9);
    HospitalSelectedImportFile.current(0);

    HospitalSelectedImportFile.bind("<<ComboboxSelected>>",HospitalChooseImportFile)
    HospitalSelectedImportFile.pack();
    HospitalSelectedImportFile.place(x=20, y=100);



    # Absorb materials list data from db
    #################Absorb materials list DB#########################
    cursor = db.cursor();
    cursor.execute("SELECT idmaterial,materialName FROM material");
    Material_in_db = cursor.fetchall();
    print(Material_in_db);

    #Create Material Drop-Dowm menu

    #matrial label
    MaterialListLabel = Label(ImportFilePage, text="Material",bg='white');
    MaterialListLabel.pack();
    MaterialListLabel.place(x=20, y=300);

    def MaterialsSelectedeImportFile(MaterialSelectedEvent):
        """Function for create Material Drop-Down menu -absorb data from DB"""
        ChoosenMaterialNewOrder=MaterialsSelectedImportFile.get();
        print(ChoosenMaterialNewOrder);
        print(ChoosenMaterialNewOrder);
        ChoosenMaterial2=ChoosenMaterialNewOrder;
        temp_Var_=list(ChoosenMaterial2);
        TempList[2]=ChoosenMaretrialIDNewOrderManual=int(temp_Var_[0]);
        print(TempList[2]);


    MaterialsSelectedImportFile = ttk.Combobox(ImportFilePage,state="readonly",value=Material_in_db,width=9);
    MaterialsSelectedImportFile.current(0);

    MaterialsSelectedImportFile.bind("<<ComboboxSelected>>",MaterialsSelectedeImportFile)
    MaterialsSelectedImportFile.pack();
    MaterialsSelectedImportFile.place(x=18, y=330);


    def SaveToDB():
        # Function to save order into DB from Import file
        global counterOrderId;
        ######################################################################################################
        def delete_label(label):#Function for Clear label
            label.destroy()


        cursor = db.cursor(buffered=True);
        print("order trying to get in DB-Add pressed");

        try:
            for i in range(1,len(AmountListFromDoc)):
                ValuseTuple=(i,TempList[1],InjectionTImeListFromdoc[i],AmountListFromDoc[i], TempList[0],TempList[2]);#BatchId=null
                cursor.execute("INSERT INTO orders (DoseNumber,Date,injection_time,amount,hospitalID,materialID) VALUES (%s,%s,%s,%s,%s,%s);",ValuseTuple);#BatchId=null
                counterOrderId=counterOrderId+1;
            updateOrdersTreeMainPageOutputOnly();##update orders tree main page
        except (mysql.connector.errors.DatabaseError,UnboundLocalError):
            logging.error(traceback.format_exc());
            DateInputCheckMsg=Label(ImportFilePage, text="Please choose date of order",fg="red", font=('Helvetica 12'));
            DateInputCheckMsg.pack();
            DateInputCheckMsg.place(x=20,y=270);
            root.after(5000, delete_label, DateInputCheckMsg); ##Clear label after 5 secondes

            HospitalInputMsg=Label(ImportFilePage, text="Please choose hospital",fg="red", font=('Helvetica 12'));
            HospitalInputMsg.pack();
            HospitalInputMsg.place(x=20,y=130);
            root.after(5000, delete_label, HospitalInputMsg);#Clear label after 5 secondes
        except Exception as e:
            logging.error(traceback.format_exc());
            #messagebox.showerror("Error message","Error !");
            print("Error");

        updateOrdersTreeMainPageOutputOnly();
        ImportFilePage.destroy();##Close import file window
        #Commit changes in DB and close connection
        db.commit()
        cursor.close()


    #Create a save button
    # saveFileIcon = Image.open("./Images/saveIcon.png");
    # save_next_Icon = saveFileIcon.resize((100,50), Image.ANTIALIAS);
    # saveImg = ImageTk.PhotoImage(save_next_Icon);
    saveButton=Button(ImportFilePage,text="Save",command=SaveToDB);
    saveButton.pack();
    saveButton.place(x=250, y=320);

    #Create a Cancel button
    # CancelIcon2 = Image.open("./Images/CancelIcon.png");
    # resized_Cancel_Icon2 = CancelIcon2.resize((100,50), Image.ANTIALIAS);
    # CancelImg2 = ImageTk.PhotoImage(resized_Cancel_Icon2);
    CancelButton2=Button(ImportFilePage,text="Cancel",command=lambda: [ImportFilePage.destroy()]);#close window-not working
    CancelButton2.pack();
    CancelButton2.place(x=450, y=320);

    #Create a File button+Label
    FileLabel = Label(ImportFilePage, text="File",bg='white');
    FileLabel.pack();
    FileLabel.place(x=500, y=65);

    FileIcon = Image.open("D:\PythonProjects\Cyclotron\Images\FileIcon.png")
    resized_File_Icon = FileIcon.resize((60,60), Image.ANTIALIAS)
    file_Image = ImageTk.PhotoImage(resized_File_Icon)
    FileButton=Button(ImportFilePage, image=file_Image, borderwidth=0,command=ImportFilefunction)
    FileButton.pack()
    FileButton.place(x=500, y=95);

    #Create a Injection Date

    InjectionDateLabel = Label(ImportFilePage, text="Injection Date",bg='white');
    InjectionDateLabel.pack();
    InjectionDateLabel.place(x=20, y=180);
    InjectionDateLabel2 = Label(ImportFilePage, text="Pick a date",fg="gray",font=("Halvetica",10));#fg=color of text
    InjectionDateLabel2.pack();
    InjectionDateLabel2.place(x=20, y=210);

    #Calender
    #add calender icon
    CalendarIcon = Image.open("D:\PythonProjects\Cyclotron\Images\CalendarIcon.png");
    resizedCalenderIcon = CalendarIcon.resize((23,23), Image.ANTIALIAS);
    CalenderImg = ImageTk.PhotoImage(resizedCalenderIcon);
    CalenderLabelicon=Label(ImportFilePage,image=CalenderImg);
    CalenderLabelicon.pack();
    CalenderLabelicon.place(x=180, y=238);

    date_today = datetime.now() # today's date

    #Add calender widget/method
    selectedDate=tk.StringVar() # declaring string variable
    def print_sel(e):
        ChoosenDateForImport=cal1.get_date();
        TempList[1]=ChoosenDateForImport;
        print( TempList[1]);

    cal1=DateEntry(ImportFilePage,background='darkblue',foreground='white',mindate=date_today,textvariable=selectedDate);
    cal1.pack(pady = 20);
    cal1.place(x=20, y=240);
    cal1.bind("<<DateEntrySelected>>", print_sel);#catch date event

    #cal1=Calendar(ImportFilePage,font="Arial 14", selectmode='day',year=2022, month=5, disableddaybackground="red", day=22);

    # for i in range(len(page_stuff)):
    #  OrdersTree.insert('', 'end',values=i)
    #OrdersTree.insert(parent='', index=0, text='', values=(TempList[0],TempList[1]));

    # def my_upd(*args): # triggered when value of string varaible changes
    #     l1.config(text=sel.get()); # read and display date
    #
    # l1=tk.Label(ImportFilePage,bg='yellow');  # Label to display date
    # l1.pack();
    #
    # sel.trace('w',my_upd) ;# on change of string variable
    # print(sel);

    #print(TempList);
    ImportFilePage.mainloop();



# iid=0
# for hospital in hospitals_in_db:
#     #print(hospital)
#     hospitals_list.insert(parent='', index='end', iid=iid, text='',
#                           values=(hospital[1], hospital[2], hospital[3]))
#     iid +=1
#
# hospitals_list.pack()


###################Buttons for edit,delete,import file and etc.###################################
#Create a button in the main Window to open the popup
# edit_button = Button(hospitalFrame, text= "Edit", command= open_popup_hospital)
# edit_button.pack(side= LEFT)
# edit_button.place(x=450, y=50)
# edit_button.pack(side=LEFT, padx=PlaceLable_X+100, pady=PlaceLable_Y+50)

#clear/delete Edit page tree-all records
# def disabledTree():
#     #"""Initialization-clear/Delete all the records"""
#     for child in NewOrderTree_P2.winfo_children():
#         child.configure(state='disable')

############################new order page###################################################

def PopUpForNewOrder():

    root = tk.Tk()
    root.title("New Order")
    root.geometry("700x600");
    root.configure(bg="#F0F3F4")

    #NewOrderMainPage =Toplevel(root);
    NewOrderMainPage =tk.Frame(root);
    NewOrderMainPage.config(bg="#F0F3F4");#color of page-white-gray


    #Create secnd pop-up window for order page
    NewOrdersecondaryPage = tk.Frame(root);
    NewOrdersecondaryPage.config(bg="#F0F3F4");#color of page-white-gray


    ########################page number 1,New order page#########################################################

    ##

    NeworderTitleLabel=Label(NewOrderMainPage, text="New Order ",bg="#F0F3F4", font=('Helvetica 17 bold'), fg='#034672');
    NeworderTitleLabel.pack();
    NeworderTitleLabel.place(x=270,y=25);
    # labels
    #Create hospital Drop-down menu

    HospitalListLabel = Label(NewOrderMainPage, text="Hospital",bg='white');
    HospitalListLabel.pack();
    HospitalListLabel.place(x=20, y=70);

    # Absorb hosital list data from db
    cursor = db.cursor();
    cursor.execute("SELECT idhospital,Name FROM hospital");
    hospitalsListForNewOrderManual = cursor.fetchall();

    def HospitalChooseNewOrder(HospitalSelectedEvent):
        """Function for create Hospital Drop-Down menu -absorb data from DB"""
        global hospitalId;
        ChoosenHospitalNewOrder=(HospitalSelectedNewOrder.get());

        #loop for find hospital name in the string
        HospitalNameFromChoosenHospital = "";
        for index in ChoosenHospitalNewOrder:
            if index.isdigit():
                pass;
            else:
                HospitalNameFromChoosenHospital = HospitalNameFromChoosenHospital + index;


        ChoosenhospitalNameFromDropDown=HospitalNameFromChoosenHospital;

        print(ChoosenhospitalNameFromDropDown);
        #Print to the screen the hospital selected-print to page number 2
        HospitalLabel2=Label(NewOrdersecondaryPage, text= ChoosenhospitalNameFromDropDown, bg="white", font=('Helvetica 14'));
        HospitalLabel2.pack();
        HospitalLabel2.place(x=20,y=80);

        #loop for findinf Id number in the string
        HospitalIDFromChoosenHospital = "";
        for index in ChoosenHospitalNewOrder:
            if index.isdigit():
                HospitalIDFromChoosenHospital = HospitalIDFromChoosenHospital + index;


        hospitalId=int(HospitalIDFromChoosenHospital);


    HospitalSelectedNewOrder = ttk.Combobox(NewOrderMainPage,state="readonly",value=hospitalsListForNewOrderManual,width=15);
    HospitalSelectedNewOrder.current(0);

    HospitalSelectedNewOrder.bind("<<ComboboxSelected>>",HospitalChooseNewOrder);
    HospitalSelectedNewOrder.pack();
    HospitalSelectedNewOrder.place(x=20, y=100);


    ###########################################################################

    # Absorb materials list data from db
    cursor = db.cursor();
    cursor.execute("SELECT idmaterial,materialName FROM material");
    Material_in_db1 = cursor.fetchall();
    print(Material_in_db1);

    #Create Material Drop-Dowm menu

    #matrial label
    MaterialListLabel1 = Label(NewOrderMainPage, text="Material",bg='white');
    MaterialListLabel1.pack();
    MaterialListLabel1.place(x=400, y=310);


    def MaterialsSelectedeNewOrder(MaterialSelectedEvent):
        global ChoosenMaretrialIDNewOrderManual;
        """Function for create Material Drop-Down menu -absorb data from DB"""
        ChoosenMaterialNewOrder=MaterialsSelectedNeworder.get();
        print(ChoosenMaterialNewOrder);
        temp_Var_=list(ChoosenMaterialNewOrder);
        ListofVal[7]=ChoosenMaretrialIDNewOrderManual=int(temp_Var_[0]);
        print(ListofVal[7]);


    MaterialsSelectedNeworder = ttk.Combobox(NewOrderMainPage,state="readonly",value=Material_in_db1,width=9);
    MaterialsSelectedNeworder.current(0);

    MaterialsSelectedNeworder.bind("<<ComboboxSelected>>",MaterialsSelectedeNewOrder);

    MaterialsSelectedNeworder.pack();
    MaterialsSelectedNeworder.place(x=400, y=330);
    global OrderID,idCounter,amount,ChoosenMaretrialIDNewOrderManual;
    OrderID=0;
    # declaring string variable for storing time interval
    TimeIntervals=tk.StringVar(NewOrderMainPage);
    # declaring string variable for storing time
    HoursVar = StringVar(NewOrderMainPage);
    MinutesVar = StringVar(NewOrderMainPage);
    amountVar=tk.StringVar(NewOrderMainPage);

    # declaring string variable for storing amount
    ListofVal=["","","","","","","",""];

    ListofTimeIntervals=[];


    def submitToNextPage():
        global amount;
        global hospitalId;
        global OrderID;
        #disabledTree();
        #Initialization-clear/Delete all the records
        for rawselected in NewOrderTree_P2.get_children():
            NewOrderTree_P2.delete(rawselected);

        def destroy_widget(widget):
            widget.destroy()


        #Get event values-events
        Time_Intervals=TimerangeLabelEntry.get();
        IntAmount=AmountOfDosesLabelEntry.get();
        Minutes_Var=MinutesCLockSelected.get();
        Hours_Var=HoursClockedSelected.get();

        #message box if not try to click next if inputs are empty-Input check
        #Check input material
        # try:
        #     ChoosenMaretrialIDNewOrderManual;
        # except NameError:
        #     MatrerialInputCheckMsg=Label(NewOrderMainPage, text="Please choose Material ",fg="red", font=('Helvetica 12'));
        #     MatrerialInputCheckMsg.pack();
        #     MatrerialInputCheckMsg.place(x=400,y=360);
        #     root.after(5000, destroy_widget, MatrerialInputCheckMsg) ##Clear label after 5 secondes
        # else:
        #     print(ChoosenMaretrialIDNewOrderManual);
        #try:
        cursor = db.cursor();
        SearchSpecOrderQueryByDoubleclick=f'SELECT idorders,CONCAT(hospitalID,Date) AS AdressPlusHospital FROM orders where hospitalID="{hospitalId}" AND Date="{ChoosenDateForManaulOrder}"';
        cursor.execute(SearchSpecOrderQueryByDoubleclick);
        OrderDatatoSpecificOrder = cursor.fetchall();

        db.commit();
        cursor.close();

        for x in OrderDatatoSpecificOrder:
            print(f"order selected:{x}");

        if OrderDatatoSpecificOrder:
            messagebox.showerror("Error message","There is allready order with that date and hospital,cant continue!");
            raise Exception("There is allready order with that date and hospital,cant continue");


        try:
            if ChoosenMaretrialIDNewOrderManual in globals():
                MaterialInputCheckMsg=Label(NewOrderMainPage, text="Please  choose Material ",fg="red", font=('Helvetica 12'));
                MaterialInputCheckMsg.pack();
                MaterialInputCheckMsg.place(x=400,y=360);
                print("bla bla")
                root.after(5000, destroy_widget, MaterialInputCheckMsg);#Clear label after 5 secondes
            else:
                print(ChoosenMaretrialIDNewOrderManual);

        except NameError:
            MatrerialInputCheckMsg=Label(NewOrderMainPage, text="Please choose Material ",fg="red", font=('Helvetica 12'));
            MatrerialInputCheckMsg.pack();
            MatrerialInputCheckMsg.place(x=400,y=360);
            root.after(5000, destroy_widget, MatrerialInputCheckMsg) ##Clear label after 5 secondes
        else:
            print(ChoosenMaretrialIDNewOrderManual);

        try:
            ChoosenDateForManaulOrder;
        except NameError:
            DateInputCheckMsg=Label(NewOrderMainPage, text="Please choose date of order",fg="red", font=('Helvetica 12'));
            DateInputCheckMsg.pack();
            DateInputCheckMsg.place(x=20,y=260);
            root.after(5000, destroy_widget, DateInputCheckMsg) ##Clear label after 5 secondes
        else:
            print(ChoosenDateForManaulOrder);

        try:
            hospitalId
        except NameError:
            HospitalInputMsg=Label(NewOrderMainPage, text="Please choose hospital",fg="red", font=('Helvetica 12'));
            HospitalInputMsg.pack();
            HospitalInputMsg.place(x=20,y=140);
            root.after(5000, destroy_widget, HospitalInputMsg);#Clear label after 5 secondes
        else:
            print(hospitalId);

            #loop for findinf if number in the string-time intervals
        for index in Time_Intervals:
            if index.isdigit()==False:
                NaNFlag=True;
            else:
                NaNFlag=False;

        if (Time_Intervals=='') or (NaNFlag==True):
            TImeIntervalMsg=Label(NewOrderMainPage, text="Please fill in Time Intervals",fg="red", font=('Helvetica 12'));
            TImeIntervalMsg.pack();
            TImeIntervalMsg.place(x=400,y=270);
            root.after(5000, destroy_widget, TImeIntervalMsg);#Clear label after 5 secondes
        else:
            print(Time_Intervals);

        if (Minutes_Var=='0') or (Hours_Var=='0'):
            BegTImeMsg=Label(NewOrderMainPage, text="Please fill in Beginning Time ",fg="red", font=('Helvetica 12'));
            BegTImeMsg.pack();
            BegTImeMsg.place(x=20,y=370);
            root.after(5000, destroy_widget, BegTImeMsg);#Clear label after 5 secondes
        else:
            print(Minutes_Var,Hours_Var);


        for index in IntAmount:
            if index.isdigit()==False:
                AmountOfrowsMsg1=Label(NewOrderMainPage, text="Amount of rows can accept only integer numbers ",fg="red", font=('Helvetica 12'));
                AmountOfrowsMsg1.pack();
                AmountOfrowsMsg1.place(x=400,y=165);
                root.after(5000, destroy_widget, AmountOfrowsMsg1);#Clear label after 5 secondes
        # else:
        #    NaNFlag=False;

        try:
            IntAmount=int(IntAmount);
        except ValueError:
            AmountOfrowsMsg2=Label(NewOrderMainPage, text="Please fill in Amount-Of-Rows for order ",fg="red", font=('Helvetica 12'));
            AmountOfrowsMsg2.pack();
            AmountOfrowsMsg2.place(x=400,y=140);
            root.after(5000, destroy_widget, AmountOfrowsMsg2);#Clear label after 5 secondes
        else:
            print(IntAmount);


        #print("Error")
        ListofVal[0]=idCounter=1;
        ListofVal[1]=(IntAmount/IntAmount);
        ListofVal[2]=int(Hours_Var);
        ListofVal[3]=int(Minutes_Var);
        TimeInjectionVar=f'{ListofVal[2]}:{ListofVal[3]}';
        print(TimeInjectionVar);
        ListofVal[4]=IntAmount;
        ListofVal[5]=hospitalId;
        ListofVal[6]=int(Time_Intervals);



        for record in range(int(IntAmount)):
            NewOrderTree_P2.insert("", "end",values=( ListofVal[0],ListofVal[1],f'{ListofVal[2]}:{ListofVal[3]}'));#ListofVal[0]=id,ListofVal[1]=amount/amount,ListofVal[2]=hours:ListofVal[3]=minutes
            #ListofVal[2]=ListofVal[2]+1;       #Hours jumps/intervals
            ListofVal[3]=ListofVal[6];         #Add minutes intervals
            ListofVal[0]= ListofVal[0]+1;
            TimeInjectionVar=f'{ListofVal[2]}:{ListofVal[3]}';
            print("Time per Dose:",TimeInjectionVar);
            ListofTimeIntervals.append(TimeInjectionVar);
            ListofVal[2]=ListofVal[2]+1;       #Hours jumps/intervals

        OrderID+=1;#counterID=counterID+1
        #print(OrderID);
        amountVar.set("");
        MinutesVar.set("");
        HoursVar.set("");
        print("List of intervals:",ListofTimeIntervals);

        deleteButton2['state'] = DISABLED;#Disable remove button
        AddRowButton['state'] = DISABLED;#Disable add dose/record button
        NewOrderTree_P2.state(('disabled',));#Disabled/gray-out new order tree
        dropDownInjectionT_M.configure(state="disabled");#Disabled/gray-out time-injection drop-down menu
        dropDownAmountM.configure(state="disabled");#Disabled/gray-out change amount drop-down menu
        #swap function for viewing New order page,frame 2,after pressing "next" """
        NewOrderMainPage.forget();
        NewOrdersecondaryPage.pack(fill='both',expand=1);

    # def clearwidgets():
    #     AmountOfrowsMsg1.destroy(),AmountOfrowsMsg2.destroy(),BegTImeMsg.destroy();
    #     TImeIntervalMsg.destroy(),HospitalInputMsg.destroy(),DateInputCheckMsg.destroy();
    #
    #
    #
    # searchEntry.bind("<Enter>",SearchComponent)#catch any key pressed and released from keyboard- event

    #Create Amount of Doses input
    AmountOfDosesLabel = Label(NewOrderMainPage, text="Amount of rows",bg="white");
    AmountOfDosesLabel.place(x=400, y=80);

    AmountOfDosesLabelEntry = Entry(NewOrderMainPage,textvariable=amountVar,font=("Halvetica",12));
    AmountOfDosesLabelEntry.config(width=7);#width of window

    # amount=AmountOfDosesLabelEntry.get();
    # print(amount);
    AmountOfDosesLabelEntry.insert(0, '');
    AmountOfDosesLabelEntry.pack();
    AmountOfDosesLabelEntry.place(x=400, y=120);


    #Create Injection time input/entry
    InjectionTimeLabel = Label(NewOrderMainPage, text="Injection Date",bg="white");
    InjectionTimeLabel.place(x=20, y=200);

    #Calender
    #add calender icon
    global  CalenderImg1
    CalendarIcon1 = Image.open("D:\PythonProjects\Cyclotron\Images\CalendarIcon.png");
    resizedCalenderIcon1 = CalendarIcon1.resize((23,23), Image.ANTIALIAS);
    CalenderImg1 = ImageTk.PhotoImage(resizedCalenderIcon1,master=NewOrderMainPage);
    CalenderLabelicon1=Label(NewOrderMainPage,image=CalenderImg1);
    CalenderLabelicon1.pack();
    CalenderLabelicon1.place(x=172, y=237);

    date_today = datetime.now() # today's date

    #Add calender widget/method
    selectDateEventManaulOrder=tk.StringVar(NewOrdersecondaryPage) # declaring string variable
    selectDateEventManaulOrder.set(' ');
    def print_sel(e):
        global ChoosenDateForManaulOrder;
        """ This function print to the tree/table """
        ChoosenDateForManaulOrder=cal.get_date();
        print(ChoosenDateForManaulOrder);
        #copy and past date event to page number 2
        dateLabel2=Label(NewOrdersecondaryPage, text= ChoosenDateForManaulOrder, bg="white", font=('Helvetica 12'));
        dateLabel2.pack();
        dateLabel2.place(x=20,y=110);

    def destroyNewOrderFunc():
        """Function for cancel button"""
        root.destroy();
        updateOrdersTreeMainPageOutputOnly();#Refresh/Update Main page
        OrdersTree.pack();         #open order main page immedaitly

    cal=DateEntry(NewOrderMainPage,background='darkblue',mindate=date_today,foreground='white',textvariable=selectDateEventManaulOrder);
    cal.pack(pady = 20);
    cal.config(width=20);#width of window
    cal.place(x=20, y=240);
    cal.bind("<<DateEntrySelected>>", print_sel);#catch date event


    #
    #Create Time range input/Entry
    TimerangeLabel = Label(NewOrderMainPage, text="Time Range/Intervals",bg="white");
    TimerangeLabel.place(x=400, y=200);

    TimerangeLabelEntry = Entry(NewOrderMainPage,textvariable=TimeIntervals,font=("Halvetica",12));
    TimerangeLabelEntry.config(width=7);#width of window
    #TimerangeLabelEntry.insert(0, '');
    TimerangeLabelEntry.pack();
    TimerangeLabelEntry.place(x=400, y=240);

    #Create Beginng time input
    AmountOfDosesLabel = Label(NewOrderMainPage, text="Beginning time",bg="white");
    AmountOfDosesLabel.place(x=20, y=300);
    # AmountOfDosesLabelEntry = Entry(NewOrderMainPage,font=("Halvetica",12));
    # AmountOfDosesLabelEntry.insert(0, '');
    # AmountOfDosesLabelEntry.pack();
    # AmountOfDosesLabelEntry.place(x=20, y=270);
    InjectionDateLabel2 = Label(NewOrderMainPage, text="Pick a time",fg="gray",font=("Halvetica",10));#fg=color of text
    InjectionDateLabel2.pack();
    InjectionDateLabel2.place(x=20, y=330);

    ################FIxed hour,minutes,and secondes########################
    #Hours
    HoursClockedSelected = Spinbox(NewOrderMainPage,textvariable=HoursVar,from_= 0, to = 24,wrap=True,width=2);
    HoursClockedSelected.pack();
    HoursClockedSelected.place(x=20, y=350);
    # Minutes
    MinutesCLockSelected = Spinbox(NewOrderMainPage,textvariable=MinutesVar ,from_= 0, to = 59,wrap=True,width=2);
    MinutesCLockSelected.pack();
    MinutesCLockSelected.place(x=65, y=350);

    # #Create a Cancel button
    # global CancelImg;
    # CancelIcon = Image.open("./Images/CancelButton.png");
    # resized_Cancel_Icon = CancelIcon.resize((100,50), Image.ANTIALIAS);
    # CancelImg = ImageTk.PhotoImage(resized_Cancel_Icon);
    CancelButtonNewOrderPage1=Button(NewOrderMainPage,text="Cancel",command=destroyNewOrderFunc);#close window-not working
    CancelButtonNewOrderPage1.pack();
    CancelButtonNewOrderPage1.place(x=400, y=530);


    #Submit/Next button
    #global NextButton;
    # nextIcon = Image.open("./Images/nextButton.png");
    # resized_next_Icon = nextIcon.resize((100,50), Image.ANTIALIAS);
    # NextButton = ImageTk.PhotoImage(resized_next_Icon);
    sub_btn=tk.Button(NewOrderMainPage ,text="Next", command = submitToNextPage)
    sub_btn.pack();
    sub_btn.place(x=200, y=530)

    def submitToNextPage1(e):
        """Enter key pressed replace next button"""
        submitToNextPage();
    root.bind('<Return>', submitToNextPage1)


    NewOrderMainPage.pack(fill='both',expand=1);


    #####################################New order 2, second page numer 2###################################################

    NewOrdersecondaryLabel=Label(NewOrdersecondaryPage, text="New Order",bg="#F0F3F4", font=('Helvetica 17 bold'), fg='#034672');
    NewOrdersecondaryLabel.pack();
    NewOrdersecondaryLabel.place(x=270,y=27);


    def enterToDB():#Function to insert data into My-SQL Db
        # ordersFrame.forget();
        # toolbar.forget();
        #ordersFrame.wm_state('iconic');#minimize orders main page
        global counterOrderId;
        MsgBox = messagebox.askquestion ('Info message','Do you wish to proceed? every changed will be saved in the DB',icon = 'warning')
        if MsgBox == 'yes':
            # ValuseDic = {
            #        'idorders': 4,
            #        'Date': '2002-03-92',
            #        'injection_time': '11:20:11',
            #        'amount': 7,
            #        'hospitalID': 7,
            #        'batchID': 7,
            #        'DecayCorrected': 7 }  ;

            #Check if all input was enterd
            if (ListofVal[7]!='') and (ListofVal[6]!='') and (ListofVal[5]!='') and (ListofVal[4]!='') and (ListofVal[3]!='') and (ListofVal[2]!='') and (ListofVal[1]!='') and (ListofVal[0]!=''):
                NewOrderTree_P2.state(('!disabled',));#Enable tree items
                deleteButton2['state'] = NORMAL;#Enable remove button
                AddRowButton['state'] = NORMAL;#Enable add dose/record button
                dropDownInjectionT_M.configure(state="normal");#Disabled/gray-out time-injection drop-down menu
                dropDownAmountM.configure(state="normal");#Disabled/gray-out change amount drop-down menu
            else:
                NewOrderTree_P2.state(('disabled',));#Disable tree items
                messagebox.showerror("Error message","Error ! one of the records not filled,please fill order again");
                root.destroy();#close New order page 2
                PopUpForNewOrder();#return to New order page


            if (Stage1forNewOrderBut['state'] == NORMAL):
                SaveOrderbutton['state'] = NORMAL;
                CancelButtonPagenumber2NewORder['state']=DISABLED;
                Stage1forNewOrderBut['state']=DISABLED;
            else:
                SaveOrderbutton['state'] = DISABLED;

            cursor = db.cursor(buffered=True);
            for i in range(1,ListofVal[4]+1):
                ValuseTuple=(i, ChoosenDateForManaulOrder, ListofTimeIntervals[i-1], ListofVal[1], ListofVal[5],ListofVal[7]);# default BatchId=null,default Decaycorrected=null
                print("order trying to get in DB-Add pressed");
                try:
                    UpdateSQlQuery="INSERT INTO orders (DoseNumber,Date,injection_time,amount,hospitalID,materialID) VALUES (%s,%s,%s,%s,%s,%s);";#BatchId=null
                    cursor.execute(UpdateSQlQuery,ValuseTuple);
                    counterOrderId=counterOrderId+1;
                    print("DB updated successfully ");
                except Exception as e:
                    logging.error(traceback.format_exc());
                    messagebox.showerror("Error message","Error !");
                    print("Error-Order was not updated-please check MySQL")

            #destroyNewOrderFunc();
            #Commit changes in DB
            db.commit()
            cursor.close()
            #Close connection to DB
            #db.close()
        else:
            messagebox.showinfo('Return','You will now return to the application screen');
            NewOrdersecondaryPage.destroy();#close New order page 2
            PopUpForNewOrder();#return to New order page




    #Create Save order/ADD button
    # global addImg;
    # AddFileIcon = Image.open("./Images/AddButton.png");
    # resized_add_Icon = AddFileIcon.resize((100,50), Image.ANTIALIAS);
    # addImg = ImageTk.PhotoImage(resized_add_Icon);
    Stage1forNewOrderBut=Button(NewOrdersecondaryPage,text="Create Order & Start edit",command=enterToDB);
    Stage1forNewOrderBut.pack();
    Stage1forNewOrderBut.place(x=170, y=520);



    # #Create a Cancel button
    # global CancelImg;
    # CancelIcon = Image.open("./Images/CancelButton.png");
    # resized_Cancel_Icon = CancelIcon.resize((100,50), Image.ANTIALIAS);
    # CancelImg = ImageTk.PhotoImage(resized_Cancel_Icon);
    CancelButtonPagenumber2NewORder=Button(NewOrdersecondaryPage,text="Cancel",command=destroyNewOrderFunc);#close window-not working
    CancelButtonPagenumber2NewORder.pack();
    CancelButtonPagenumber2NewORder.place(x=450, y=520);


    #Empty page/table for new order,create New tree for page 2
    NewOrderTree_P2 = ttk.Treeview(NewOrdersecondaryPage,height=15);
    NewOrderTree_P2.pack();
    NewOrderTree_P2.place(x=170,y=130);

    #New order page 2 page/tree scrollbar -vertical
    NewOrderTree_P2_Scroollbar = Scrollbar(NewOrdersecondaryPage, orient="vertical", command=NewOrderTree_P2.yview);
    NewOrderTree_P2_Scroollbar.place(x=495, y=130, height=330)
    NewOrderTree_P2.configure(yscrollcommand = NewOrderTree_P2_Scroollbar.set);

    # Defining number of columns
    NewOrderTree_P2['columns']= ("ID","Amount","Injection time");
    #Foramte Columns
    NewOrderTree_P2.column("#0",width=0,minwidth=0);
    NewOrderTree_P2.column("ID",anchor=W,width=80,minwidth=25);
    NewOrderTree_P2.column("Amount",anchor=CENTER,width=120,minwidth=25);
    NewOrderTree_P2.column("Injection time",anchor=W,width=120,minwidth=25);

    #Define headers/titles in table
    NewOrderTree_P2.heading("#0", text="Label",anchor=W);
    NewOrderTree_P2.heading("ID", text="Dose Number",anchor=W);
    NewOrderTree_P2.heading("Amount", text="Amount",anchor=CENTER);
    NewOrderTree_P2.heading("Injection time", text="Injection time",anchor=W);

    #############################Amount quantity Event ######################

    AmountList = [
        "0","1","2","3","4","5"
    ]

    status = tk.StringVar()
    status.set("0")

    #Catch event
    def AmountSelectedForNewOrder(event):
        row = NewOrderTree_P2.focus();
        dataofchoosnenRowListEditTree=row;
        print(dataofchoosnenRowListEditTree);
        DataOfRowSelectedDicEditTree=NewOrderTree_P2.item(dataofchoosnenRowListEditTree);
        DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        print("Raw selectes to edit:",DataOfRowSelectedList);
        IidSelected=DataOfRowSelectedList[0];
        AmountSelected=DataOfRowSelectedList[1];
        InjectionTimeSelected=DataOfRowSelectedList[2];

        if row:
            status.set(NewOrderTree_P2.set(row, 'two'))

    NewOrderTree_P2.bind('<<TreeviewSelect>>', AmountSelectedForNewOrder)

    def set_amountValueEventNewOrderP(CurAmountvalue):
        row = NewOrderTree_P2.focus();
        print("Amount selected",CurAmountvalue);
        if row:
            try:
                cursor = db.cursor(buffered=True);
                UpdateSQlQuery=f"UPDATE  orders SET amount='{CurAmountvalue}'  WHERE DoseNumber = '{DoseNum_Selected}' AND hospitalID= '{ListofVal[5]}' AND Date= '{ChoosenDateForManaulOrder}';";
                cursor.execute(UpdateSQlQuery);
                print("DB updated(Amount) successfully ");
                db.commit();
                cursor.close();
            except Exception as e:
                logging.error(traceback.format_exc())
                print("Error-Order was not updated-please check MySQL")
            NewOrderTree_P2.set(row, '1', CurAmountvalue)


    dropDownAmountM = ttk.OptionMenu(NewOrdersecondaryPage, status, "0", *AmountList, command=set_amountValueEventNewOrderP);
    dropDownAmountM.pack();
    dropDownAmountM.place(x=250, y=490);
    #Change Amount time manual label
    ChangeAmountLabel=Label(NewOrdersecondaryPage, text="Change Amount : ", font=('Helvetica 12'));
    ChangeAmountLabel.pack();
    ChangeAmountLabel.place(x=100,y=490);

    #############################Amount Event over######################

    # ############################Injection Time event#########################
    TimeList = [
        "06:00","06:30","07:00","07:30","08:00","08:30","09:00","09:30","10:00","10:30","11:00",
        "11:30","12:00","12:30","13:00","13:30","14:00","14:30","15:00","15:30","16:00","16:30",
        "17:00","17:30","18:00","18:30","19:00","19:30","20:00"
    ]

    status = tk.StringVar()
    status.set("00:00")

    #Catch Injection  event
    def InjectionTimeselect(event):
        row = NewOrderTree_P2.focus();
        global DoseNum_Selected,AmountSelected,InjectionTimeSelected,DosesSelectedEvent;
        row = NewOrderTree_P2.focus();

        dataofchoosnenRowListEditTree=row;
        print(dataofchoosnenRowListEditTree);
        DataOfRowSelectedDicEditTree=NewOrderTree_P2.item(dataofchoosnenRowListEditTree);
        DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        print("Raw selectes to edit 2:",DataOfRowSelectedList);
        DoseNum_Selected=DataOfRowSelectedList[0];
        AmountSelected=DataOfRowSelectedList[1];
        InjectionTimeSelected=DataOfRowSelectedList[2];
        if row:
            status.set(NewOrderTree_P2.set(row, 'Injection time'))

    NewOrderTree_P2.bind('<<TreeviewSelect>>', InjectionTimeselect)

    def setInjectionTime(CurTimevalueEditNewOrderP):
        row = NewOrderTree_P2.focus()
        if row:
            try:
                cursor = db.cursor(buffered=True);
                UpdateSQlQuery=f"UPDATE  orders SET injection_time='{CurTimevalueEditNewOrderP}'  WHERE DoseNumber = '{DoseNum_Selected}' AND hospitalID= '{ListofVal[5]}' AND Date= '{ChoosenDateForManaulOrder}';";
                cursor.execute(UpdateSQlQuery);
                print("DB updated(Injection_Time) successfully ");
                db.commit();
                cursor.close();
            except Exception as e:
                logging.error(traceback.format_exc())
                print("Error-Order was not updated-please check MySQL")
            NewOrderTree_P2.set(row, '2', CurTimevalueEditNewOrderP)


    dropDownInjectionT_M = ttk.OptionMenu(NewOrdersecondaryPage, status, "00:00", *TimeList, command=setInjectionTime);
    dropDownInjectionT_M.pack();
    dropDownInjectionT_M.place(x=300, y=460);
    #Change injecion time manual label
    ChangeTImeIjectionLabel=Label(NewOrdersecondaryPage, text="Change Time-Injection : ", font=('Helvetica 12'));
    ChangeTImeIjectionLabel.pack();
    ChangeTImeIjectionLabel.place(x=100,y=460);


    def enterToDBAfterEdit():
        #Function to insert data into My-SQL Db
        destroyNewOrderFunc();
        # root.destroy();#Close import file-manual window
        updateOrdersTreeMainPageOutputOnly();#Refresh/Update Main page
        # OrdersTree.pack();         #open order main page immedaitly




    # #Create Save order/ADD button
    # global addImg;
    # AddFileIcon = Image.open("./Images/AddButton.png");
    # resized_add_Icon = AddFileIcon.resize((100,50), Image.ANTIALIAS);
    # addImg = ImageTk.PhotoImage(resized_add_Icon);
    SaveOrderbutton=Button(NewOrdersecondaryPage,text="Save & Finish Editing",command=enterToDBAfterEdit,state = DISABLED);
    SaveOrderbutton.pack();
    SaveOrderbutton.place(x=320, y=520);


    #Create ADD row button+icon
    # defining a function that will
    # print them on the screen
    #rowTree = StringVar();
    def addRowFunc():
        #global idCounter;
        rowTreetoAdd=(ListofVal[0],ListofVal[1],ListofVal[2]);
        NewOrderTree_P2.insert("", "end", values=rowTreetoAdd);
        ListofVal[0]=ListofVal[0]+1;
        ListofVal[4]+=1;#current amount= courrent amount+1

    def removeRawFunc():
        #rowTree=rowTree.get();
        #for i,j in zip(range(IntAmount),range(BeginigHour,IntAmount)):
        rawSelectedToDelete=NewOrderTree_P2.selection();
        for rawselected in rawSelectedToDelete:
            NewOrderTree_P2.delete(rawselected);
        ListofVal[4]=ListofVal[4]-1;#current amount= current amount-1
    #amountVar.set("");

    #Remove button (Icon) - List-ORders page number 2
    global imgDelete2;
    deleteIcon2 = Image.open("D:\PythonProjects\Cyclotron\‏‏deleteIcon.png");
    resizedDeleteIcon2 = deleteIcon2.resize((25,25), Image.ANTIALIAS);
    imgDelete2 = ImageTk.PhotoImage(resizedDeleteIcon2,master=NewOrdersecondaryPage);
    deleteButton2=Button(NewOrdersecondaryPage,image=imgDelete2,bg="white",font=('Helvetica 14'), borderwidth=0,command=removeRawFunc);
    deleteButton2.pack();
    deleteButton2.place(x=465, y=95);




    ####################Buttons for new order-manual page##########################


    global addROWImg;
    AddrowLabel=Label(NewOrdersecondaryPage, text="Add row",bg="white", font=('Helvetica 14'));
    AddrowLabel.pack();
    AddrowLabel.place(x=270,y=98);
    #Add row image+button
    AddrowIcon = Image.open("D:/PythonProjects/Cyclotron/addIcon.png");
    resized_add_Row = AddrowIcon.resize((25,25), Image.ANTIALIAS);
    addROWImg = ImageTk.PhotoImage(resized_add_Row,master=NewOrdersecondaryPage);
    AddRowButton=Button(NewOrdersecondaryPage,image=addROWImg, borderwidth=0,command=addRowFunc);
    AddRowButton.pack();
    AddRowButton.place(x=240, y=100);

# ####################end of page number 2 -New order #######################################################################

################################Edit/Update Order main page###################################################################
def UpdateOrder(event):
    global hospitalId;
    curItem = OrdersTree.focus();
    DataOfRowSelectedDic=OrdersTree.item(curItem);
    DataOfRowSelectedList=DataOfRowSelectedDic['values'];
    #print(DataOfRowSelectedList);
    HospitalSelectedName=DataOfRowSelectedList[0];
    DateSelected=DataOfRowSelectedList[1];
    AmountOfDosesSelected=DataOfRowSelectedList[2];



    #search hospital by name from hospital table db and get the ID and the name as output
    cursor = db.cursor();
    cursor.execute(f'SELECT idhospital,Name FROM hospital where Name="{HospitalSelectedName}"');
    hospitalsListForNewOrderManual = cursor.fetchall();

    HospitalListNewOrderPage = hospitalsListForNewOrderManual;
    print(HospitalListNewOrderPage);


    EditPage =Toplevel(root);
    EditPage.title("Edit Order");
    EditPage.geometry("700x600");
    EditPage.config(bg="#F0F3F4");#Color of page(White-Gray)

    #NewOrderMainPage.place(x=450,y=70);

    #NewOrdersecondaryPage = tk.Frame(root);


    HospitalLabelForEditPage=Label(EditPage, text=HospitalSelectedName,bg="white", font=('Helvetica 14'));
    x=str(HospitalListNewOrderPage);#string type here
    print(x);
    hospitalLabel = x.split(",",1);
    hospitalNameTemp=hospitalLabel[1];
    #print(hospitalNameTemp);
    hospitalIDTemp=hospitalLabel[0];
    #print(hospitalIDTemp);
    HospitalID=hospitalIDTemp.split("(");
    hospitalId=int(HospitalID[1]);

    HospitalLabelForEditPage.pack();
    HospitalLabelForEditPage.place(x=20,y=80);

    FindActivityByHospitalID=f'SELECT Fixed_activity_level FROM hospital where idhospital="{hospitalId}"';
    cursor.execute(FindActivityByHospitalID);
    ActivitiyBySpecificOrder = cursor.fetchall();
    TempID=[i[0] for i in ActivitiyBySpecificOrder];#find index number in a list of tuple
    ActivitiyBySpecificOrder2=TempID[0];
    print("Activity of Hospital selected:",ActivitiyBySpecificOrder2);


    #Create tree/table for the Edit page
    EditTree = ttk.Treeview(EditPage,height=15,selectmode="browse");#select=browse means can choose only 1 record at the time
    EditTree['columns']= ("ID","Amount","Injection time")
    EditTree.pack();
    EditTree.place(x=170,y=130);

    #Edit page/tree scrollbar -vertical
    EditOrderPageScroll = Scrollbar(EditPage, orient="vertical", command=EditTree.yview);
    EditOrderPageScroll.place(x=515, y=130, height=330)
    EditTree.configure(yscrollcommand = EditOrderPageScroll.set);


    #Foramte Columns
    EditTree.column("#0",width=0,minwidth=0);
    EditTree.column("ID",anchor=W,width=80,minwidth=25);
    EditTree.column("Amount",anchor=CENTER,width=140,minwidth=25);
    EditTree.column("Injection time",anchor=W,width=120,minwidth=25);

    #Define headers/titles in table
    EditTree.heading("#0", text="Label",anchor=W);
    EditTree.heading("ID", text="Dose number",anchor=W);
    EditTree.heading("Amount", text=f"Amount of Doses ({ActivitiyBySpecificOrder2})",anchor=CENTER);
    EditTree.heading("Injection time", text="Injection time",anchor=W);

    #get order detail from DB by hospitalID and Date
    cursor = db.cursor();
    FindActivityByHospitalID=f'SELECT Fixed_activity_level FROM hospital where idhospital="{hospitalId}"';
    cursor.execute(FindActivityByHospitalID);
    ActivitiyBySpecificOrder = cursor.fetchall();
    TempID=[i[0] for i in ActivitiyBySpecificOrder];#find index number in a list of tuple
    ActivitiyBySpecificOrder2=TempID[0];
    print("Activity of Hospital selected:",ActivitiyBySpecificOrder2);


    SearchSpecOrderQueryByDoubleclick=f'SELECT DoseNumber,amount,Injection_time FROM orders where hospitalID="{hospitalId}" AND Date="{DateSelected}"';
    cursor.execute(SearchSpecOrderQueryByDoubleclick);
    OrderDatatoSpecificOrder = cursor.fetchall();
    print(f"order selected: {OrderDatatoSpecificOrder}");
    ListOfInjectionTime=[];
    ListOfAmount=[];
    IdOrderPK=[];
    #output orders main data from DB to the orders tree
    for record in OrderDatatoSpecificOrder:
        EditTree.insert(parent='', index='end',values=(record[0],record[1],record[2]));#record[0]=Id,record[1]=amount,record[2]=injection time
        ListOfInjectionTime.append(record[2]);
        ListOfAmount.append(record[1]);
        IdOrderPK.append((record[0]));
    #EditTree.pack();
    print(ListOfInjectionTime);
    print(ListOfAmount);
    # declaring string variable for storing amount
    amountVar=AmountOfDosesSelected;
    # declaring string variable for storing time interval
    TimeIntervals="30";
    # declaring string variable for storing time
    # HoursVar = "1";
    # MinutesVar = "1";
    global OrderID,DosesSelectedEvent,idCounter;
    #global DosesSelectedEvent;
    OrderID=0;
    #global idCounter;
    ListofVal=["","","","","","",""];
    #global hospitalId;
    #global OrderID;


    #Get Time varibles avent,hous and minutes
    Time_Intervals=TimeIntervals;
    # Minutes_Var=MinutesVar;
    # Hours_Var=HoursVar;
    #get amount event variable
    #message box if not try to click next if inputs are empty
    amount=amountVar;
    IntAmount=(int(amount));
    ListofVal[0]=idCounter=1;
    ListofVal[1]=(IntAmount / IntAmount);
    # ListofVal[2]=int(Hours_Var);
    # ListofVal[3]=int(Minutes_Var);
    ListofVal[4]=IntAmount;
    ListofVal[5]=hospitalId;
    ListofVal[6]=int(Time_Intervals);

    OrderID+=1;#counterID=counterID+1
    #print(OrderID);


    #Update Date in GUI
    ChoosenDateForManaulOrder=DateSelected;
    print(ChoosenDateForManaulOrder);
    #copy and past date event to page number 2
    dateLabel2=Label(EditPage, text= ChoosenDateForManaulOrder, bg="white", font=('Helvetica 14'));
    dateLabel2.pack();
    dateLabel2.place(x=20,y=110);


    #
    # #Create Time range input/Entry
    # TimerangeLabel = Label(NewOrderMainPage, text="Time Range/Intervals",bg="white");
    # TimerangeLabel.place(x=400, y=200);
    # TimerangeLabelEntry = Entry(NewOrderMainPage,textvariable=TimeIntervals,font=("Halvetica",12));
    # TimerangeLabelEntry.config(width=7);#width of window
    # TimerangeLabelEntry.insert(0, '');
    # TimerangeLabelEntry.pack();
    # TimerangeLabelEntry.place(x=400, y=240);

    #Title of page
    NewOrdersecondaryLabel=Label(EditPage, text="Edit order",bg="#F0F3F4", font=('Helvetica 18 bold'), fg='#034672');
    NewOrdersecondaryLabel.pack();
    NewOrdersecondaryLabel.place(x=270,y=27);




    # #Create a Cancel button
    # global CancelImg;
    # CancelIcon = Image.open("./Images/CancelButton.png");
    # resized_Cancel_Icon = CancelIcon.resize((100,50), Image.ANTIALIAS);
    # CancelImg = ImageTk.PhotoImage(resized_Cancel_Icon);
    CancelButton2=Button(EditPage,text="Done editing",command=lambda: [EditPage.destroy()]);#close window-not working
    CancelButton2.pack();
    CancelButton2.place(x=390, y=520);

    #Update values,amount and injection time####################################################
    global ClickingRowsCounter;
    ClickingRowsCounter=0;
    TempListForUpdateOrderValues=["",""];#List for saving

    #############################Amount quantity Event ######################

    AmountList = [
        "0","1","2","3","4","5"
    ]

    status = tk.StringVar()
    status.set("0")

    #Catch event
    def treeAmountSelect(event):
        global ClickingRowsCounter;

        row = EditTree.focus();#catch row item

        #Count numbers of chnages in order editing
        ClickingRowsCounter=ClickingRowsCounter+1;
        if ClickingRowsCounter>1:
            messagebox.showinfo(title="Info message", message="You cant edit more then 1 record at the time");

        #Enble  Injection-time and amount dropdown menues
        dropDownInjectionT_M.configure(state="normal ");
        dropDownAmountM.configure(state="normal ");

        dataofchoosnenRowListEditTree=row;
        print(dataofchoosnenRowListEditTree);
        DataOfRowSelectedDicEditTree=EditTree.item(dataofchoosnenRowListEditTree);
        DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        print(DataOfRowSelectedList);
        IidSelected=DataOfRowSelectedList[0];
        AmountSelected=DataOfRowSelectedList[1];
        InjectionTimeSelected=DataOfRowSelectedList[2];

        #print(row)
        if row:
            status.set(EditTree.set(row, 'Amount'));

    EditTree.bind('<<TreeviewSelect>>', treeAmountSelect);

    def set_amount(AmountValue):
        #global DosesSelectedEvent;
        print(AmountValue);
        TempListForUpdateOrderValues[1]=int(AmountValue);
        #print(TempListForUpdateOrderValues[1]);
        #DosesSelectedEvent=AmountValue;
        row = EditTree.focus();

        if row:
            #try:
            #     cursor = db.cursor(buffered=True);
            #     UpdateSQlQuery=f"UPDATE  orders SET injection_time='{InjectionTimeSelected}',amount='{AmountValue}',batchID='{0}',DecayCorrected='{0}'  WHERE idorders = '{IidSelected}';";
            #     cursor.execute(UpdateSQlQuery);
            #     print("DB updated successfully ");
            #     db.commit();
            #     cursor.close();
            # except Exception as e:
            #     logging.error(traceback.format_exc())
            #     print("Error-Order was not updated-please check MySQL")
            EditTree.set(row, '1', AmountValue);


    dropDownAmountM = ttk.OptionMenu(EditPage, status, "0", *AmountList, command=set_amount);
    dropDownAmountM.pack();
    dropDownAmountM.place(x=200, y=490);
    #Change Amount time manual label
    ChangeAmountLabel=Label(EditPage, text="Change Amount : ", font=('Helvetica 12'));
    ChangeAmountLabel.pack();
    ChangeAmountLabel.place(x=20,y=490);




    # ############################Injection Time event#########################
    TimeList = [
        "06:00","06:30","07:00","07:30","08:00","08:30","09:00","09:30","10:00","10:30","11:00",
        "11:30","12:00","12:30","13:00","13:30","14:00","14:30","15:00","15:30","16:00","16:30",
        "17:00","17:30","18:00","18:30","19:00","19:30","20:00"
    ]

    Timeselected = tk.StringVar()
    Timeselected.set("00:00")
    #print(status);
    #Catch Injection  event
    def InjectionTimeselect(event):
        global DoseNum_Selected,AmountSelected,InjectionTimeSelected,DosesSelectedEvent;

        global ClickingRowsCounter;

        row = EditTree.focus();
        ClickingRowsCounter=ClickingRowsCounter+1;
        if ClickingRowsCounter>1:
            messagebox.showinfo(title="Info message", message="You cant edit more then 1 record at the time");


        #Enble  Injection-time and amount dropdown menues
        dropDownInjectionT_M.configure(state="normal ");
        dropDownAmountM.configure(state="normal ");

        dataofchoosnenRowListEditTree=row;
        #print("Row selected",dataofchoosnenRowListEditTree);
        DataOfRowSelectedDicEditTree=EditTree.item(dataofchoosnenRowListEditTree);
        DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        print("Row/Values selected",DataOfRowSelectedList);
        DoseNum_Selected=DataOfRowSelectedList[0];
        AmountSelected=DataOfRowSelectedList[1];
        TempListForUpdateOrderValues[1]=AmountSelected;#insert current amount selected to the list-Defualt
        InjectionTimeSelected=DataOfRowSelectedList[2];
        TempListForUpdateOrderValues[0]=InjectionTimeSelected;#insert current time_injection selected to the list-Defualt

        if row:
            Timeselected.set(EditTree.set(row, 'Injection time'));

    EditTree.bind('<<TreeviewSelect>>', InjectionTimeselect)

    def setInjectionTime(CurValueForTime):
        global DoseNum_Selected,AmountSelected,InjectionTimeSelected,DosesSelectedEvent;
        print(CurValueForTime);
        TempListForUpdateOrderValues[0]=CurValueForTime;
        #print(f'Injection time selected:{TempListForUpdateOrderValues[0]}');
        row = EditTree.focus();
        # dataofchoosnenRowListEditTree=row;
        # #print(dataofchoosnenRowListEditTree);
        # DataOfRowSelectedDicEditTree=EditTree.item(dataofchoosnenRowListEditTree);
        # DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        # print(DataOfRowSelectedList);
        if row:
            EditTree.set(row, '2', CurValueForTime)

    dropDownInjectionT_M = ttk.OptionMenu(EditPage, Timeselected, "00:00", *TimeList, command=setInjectionTime);
    dropDownInjectionT_M.pack();
    dropDownInjectionT_M.place(x=200, y=460);
    #Change injecion time manual label
    ChangeTImeIjectionLabel=Label(EditPage, text="Change Time-Injection : ", font=('Helvetica 12'));
    ChangeTImeIjectionLabel.pack();
    ChangeTImeIjectionLabel.place(x=20,y=460);


    def enterToDB_UpdateOrder():
        """Function for submitting/save changes from edit order page"""
        global DoseNum_Selected,ClickingRowsCounter;
        ClickingRowsCounter=0;
        #Disable  Injection-time and amount dropdown menues
        dropDownInjectionT_M.configure(state="disabled")
        dropDownAmountM.configure(state="disabled")
        #print(TempListForUpdateOrderValues[0]);

        try:
            cursor = db.cursor(buffered=True);
            UpdateSQlQuery=f"UPDATE  orders SET injection_time='{TempListForUpdateOrderValues[0]}',amount='{TempListForUpdateOrderValues[1]}' WHERE hospitalID = '{hospitalId}' and Date = '{DateSelected}' and DoseNumber={DoseNum_Selected};";
            cursor.execute(UpdateSQlQuery);
            print("DB updated successfully ");
            db.commit();
            cursor.close();
        except Exception as e:
            logging.error(traceback.format_exc())
            messagebox.showerror("Error message","Error !");
            print("Error-Order was not updated-please check MySQL")


        #Function to insert data into My-SQL Db
        #EditPage.destroy();#Close import file-manual window
        updateOrdersTreeMainPageOutputOnly();#Refresh/Update Main page
        #OrdersTree.pack();         #open order main page immedaitly
    #

    #Create ADD button
    # global addImg;
    # AddFileIcon = Image.open("./Images/AddButton.png");
    # resized_add_Icon = AddFileIcon.resize((100,50), Image.ANTIALIAS);
    # addImg = ImageTk.PhotoImage(resized_add_Icon);
    #
    AddButton1=Button(EditPage,text="Update record",command=enterToDB_UpdateOrder);
    AddButton1.pack();
    AddButton1.place(x=180, y=520);

    #Disable  Injection-time and amount dropdown menues
    dropDownInjectionT_M.configure(state="normal")
    dropDownAmountM.configure(state="normal")




    #Create ADD row button+icon


    # defining a function that will
    # print them on the screen
    #rowTree = StringVar();
    def addRowFunc():
        #global idCounter;
        rowTreetoAdd=(ListofVal[0],ListofVal[1],ListofVal[2]);
        EditTree.insert("", "end", values=rowTreetoAdd);
        ListofVal[0]=ListofVal[0]+1;
        ListofVal[4]+=1;#current amount= courrent amount+1

    def removeRawFunc():
        global DoseNum_Selected;
        #rowTree=rowTree.get();
        #for i,j in zip(range(IntAmount),range(BeginigHour,IntAmount)):
        rawSelectedToDelete=EditTree.selection();
        for rawselected in rawSelectedToDelete:
            try:
                EditTree.delete(rawselected);
                cursor = db.cursor(buffered=True);
                DeleteRecordOrdersTable_query = f'DELETE FROM orders WHERE idorders = {DoseNum_Selected}';
                cursor.execute(DeleteRecordOrdersTable_query);
                print("Remove record from DB sucssful ");
                db.commit();
                cursor.close();
            except Exception as e:
                logging.error(traceback.format_exc())
                print("Error with,record not delete, check MySQL");

        ListofVal[4]=ListofVal[4]-1;#current amount= current amount-1
    #amountVar.set("");




    ####################Buttons for new order-manual page##########################
    # Remove button (Icon) - List
    global imgDelete2;
    deleteIcon2 = Image.open("D:\PythonProjects\Cyclotron\‏‏deleteIcon.png");
    resizedDeleteIcon2 = deleteIcon2.resize((25,25), Image.ANTIALIAS);
    imgDelete2 = ImageTk.PhotoImage(resizedDeleteIcon2);
    deleteButton2=Button(EditPage, image=imgDelete2, borderwidth=0,command=removeRawFunc);
    deleteButton2.pack();
    deleteButton2.place(x=460, y=98);


    #Creaet Label Add row
    global addROWImg;
    AddrowLabel=Label(EditPage, text="Add row",bg="white", font=('Helvetica 14'));
    AddrowLabel.pack();
    AddrowLabel.place(x=280,y=97);
    #Add row image+button
    AddrowIcon = Image.open("D:/PythonProjects/Cyclotron/addIcon.png");
    resized_add_Row = AddrowIcon.resize((25,25), Image.ANTIALIAS);
    addROWImg = ImageTk.PhotoImage(resized_add_Row);
    AddRowButton=Button(EditPage,image=addROWImg, borderwidth=0,command=addRowFunc);
    AddRowButton.pack();
    AddRowButton.place(x=250, y=98);

#Double click on  main order page tree event
OrdersTree.bind('<<TreeviewOpen>>', UpdateOrder)

###############################End of update page#################################################

###########################Spacial buttons#################################################

#Create a button for import orders files (Excel or Word)
# ImportFileIcon = Image.open("ImportFile2.png")
# ImportFileIcon = Image.open("ImportFile2.png")
# resized_Edit_Icon = ImportFileIcon.resize((80,20), Image.ANTIALIAS)
# img_Edit = ImageTk.PhotoImage(resized_Edit_Icon)
importFileButton=Button(ordersFrame, text="Import file",command=importFileFunc)
importFileButton.pack()
importFileButton.place(x=250, y=65)



#Create New order button
# NewOrderIcon = Image.open("./Images/AddnewOrder2.png")
# resizedNewOrderIconIcon = NewOrderIcon.resize((120,20), Image.ANTIALIAS)
# NewOrderIconimg = ImageTk.PhotoImage(resizedNewOrderIconIcon)
editButton=Button(ordersFrame, text="Add new order",command=PopUpForNewOrder);
editButton.pack();
editButton.place(x=100, y=65);

##############end of Orders page#########################################

def Orders_page():
    """ this function is swap function for Orders frame/page"""
    workPlanButton.config(bg='#F0F0F0');  ##F0F0F0 is default color(gray)
    mbtn.config(bg='#F0F0F0'); ##F0F0F0 is default color(gray)
    ordersButton.config(bg="gray");
    ordersFrame.pack(fill='both',expand=1);
    WorkPlanFrame.forget();
    hospitalFrame.forget();
    moduleSettingsFrame.forget()
    materialSettingsFrame.forget()
    hospitalFrame.forget()
    cycloSettingsFrame.forget()

#    SettingsFrame.forget();

###########################################################################################
#start Einav code-Settings & Workplane&Batches

#general
label_font = ('Helvetica',26, 'bold')
label_font_flag_on_page = ('Helvetica 12 bold underline')
label_font_flag = ('Helvetica 12')
sub_label_font = ('Helvetica',18, 'bold')
label_color = '#034672'
red_color =  '#f5bfbf'



class table(ttk.Treeview):
    def  __init__(self,frame,scroll_width,list_height,side,x_crol,y_crol,lable_place_x,
                  lable_place_y):
        self.side = side
        scroll = Scrollbar(frame, orient="vertical", width=scroll_width)
        scroll.pack(side=side)
        scroll.place(x=x_crol, y=y_crol)
        ttk.Treeview.__init__(self,frame, yscrollcommand=scroll.set, height=list_height)
        self.pack(side=LEFT, padx=lable_place_x + 30, pady=lable_place_y + 50)
        scroll.config(command=self.yview)

        # list = self.(frame, yscrollcommand=scroll.set, height=list_height)

    # def create_fully_tabel(self,scroll_width,side, x_crol,y_crol, frame, list_height, lable_place_x,lable_place_y, columns_name_list, query):
    def create_fully_tabel(self, columns_name_list, query):

        # scroll = Scrollbar(frame, orient="vertical", width=scroll_width)
        # scroll.pack(side=side)
        # scroll.place(x=x_crol, y=y_crol)
        #
        # list = ttk.Treeview(frame, yscrollcommand=scroll.set, height=list_height)

        self['columns'] = columns_name_list

        self.column("#0", width=0, stretch=NO)
        self.heading("#0", text="", anchor=CENTER)

        i=0
        len_of_col=len(columns_name_list)
        for column_name in columns_name_list:
            # column format
            if i == 0 or i == len_of_col-2:
                width = len(column_name)*6 +30
            else:
                width = len(column_name)*6

            self.column(column_name, anchor=CENTER, width=width)
            # # Create Headings
            self.heading(column_name, text=column_name, anchor=CENTER)
        # query = query + " WHERE ISNULL(deleted)"
        cursor.execute(query)
        data = cursor.fetchall()
        # print(query,data)
        iid=0
        for recorf in data:
            val=[]
            for i in range (0,len_of_col): # plus 1 is for the pk that will not show in the table
                val.append(recorf[i+1])
            val.append(recorf[0])

            self.insert(parent='', index='end', iid=iid, text='',
                        values=val)
            iid +=1

            self.pack()

    def selected(self):
        selected = self.focus()
        selected_record = self.item(selected, 'values')
        return selected_record

    def selected_is_non(self, selected_record):
        if selected_record =='':
            text = "Please select a record from the table"
            error_message(text)
            return True
        else:
            return False

    def fk_rec_is_exist(self,query,table_name, pk_delected_record ):
        fk_list = [rec for rec in fk if rec[2]== table_name]
        if len(fk_list) != 0:
            for fk_rec in fk_list:
                query = "select * from "+ fk_rec[0] + " where " + fk_rec[1] + "=" + pk_delected_record
                try:
                    cursor.execute(query)
                    data = cursor.fetchall()
                    if data != []:
                        return True
                except:
                    #Rollback in case there is any error
                    db.rollback()
        return False

    def delete_record(self, query,table_name):
        selected_rec = self.selected()
        item_in_string= ', '.join([ item for item in selected_rec[:selected_rec.__len__()-1]])
        is_non=self.selected_is_non(selected_rec)
        if not is_non:
            len = selected_rec.__len__()
            pk_delected_record = selected_rec[len-1]
            title_tab = "Delete Record"
            text_mess= "Are you sure you want to delete " + item_in_string + " ?"
            if YES_NO_message(title_tab, text_mess):
                pk_delected_record_list = (pk_delected_record, )
                to_delete= not self.fk_rec_is_exist(query,table_name,pk_delected_record)
                pk_name = [pk[1] for pk in table_pk_list if pk[0] == table_name][0]
                if to_delete:
                    try:
                        query = "delete from "+table_name + " where " + pk_name + "=" + pk_delected_record
                        cursor.execute(query )
                        db.commit()
                    except:
                        # Rollback in case there is any error
                        db.rollback()

                else: #to hide
                    query2 = "UPDATE " + table_name +" SET deleted = True " +"WHERE " + pk_name + "=" + pk_delected_record
                    cursor.execute(query2)
                    db.commit()
                self.delete(self.selection()[0])


##################### settings new - cyclotron #####################
#cyclotron frame
cycloSettingsFrame = Frame(root)
# h = Scrollbar(cycloSettingsFrame, orient='horizontal')
# cycloSettingsFrame.pack(fill=X)

# feed label - cyclotron
feedLabel = Label(cycloSettingsFrame, text = 'Settings ➝ ', font=label_font_flag,fg=label_color)
PlaceLable_X=50
PlaceLable_Y=10
feedLabel.pack(side=LEFT)
feedLabel.place(x=PlaceLable_X,y=PlaceLable_Y)

feedLabeflag = Label(cycloSettingsFrame, text = 'Cyclotron', font=label_font_flag_on_page,fg=label_color)

PlaceLable2_X=135
feedLabeflag.pack(side=LEFT)
feedLabeflag.place(x=PlaceLable2_X,y=PlaceLable_Y)

##################### Cyclotron #####################
# Cyclotron Details label
CyclotronLabel = Label(cycloSettingsFrame, text = 'Cyclotron Details', font=sub_label_font,fg=label_color)
Lable_place_x=80
Lable_place_y=60

CyclotronLabel.pack(side=LEFT)
CyclotronLabel.place(x=Lable_place_x,y=Lable_place_y)

###cycortion tabel###
scroll_width=20
tab_side=LEFT
x=613
y= 140
frame=cycloSettingsFrame
list_height=5
table_place_x = 80
table_place_y = 80
columns_name_list=('Version', 'Capacity (mci/h)', 'Constant Efficiency (mCi/mA)', 'Description')

query = "SELECT * FROM resourcecyclotron WHERE ISNULL(deleted)"

# cyclo_tabel=table(scroll_width,tab_side, x,y,frame,list_height,lable_place_x,lable_place_y, columns_name_list, query )
cyclo_tabel=table(frame,scroll_width,list_height,tab_side,x,y,table_place_x,
                  table_place_y,)
cyclo_tabel.create_fully_tabel( columns_name_list, query)


###cycortion functions###
def editCyclotronfun():
    selected_rec = cyclo_tabel.selected()
    selected_non=cyclo_tabel.selected_is_non(selected_rec)
    if not selected_non:
        editCyclPopup = Popup()
        popup_size = "900x550"
        editCyclPopup.open_pop('Edit Cyclotron Details',popup_size)

        query = "UPDATE resourcecyclotron SET version = %s ,capacity= %s, constant_efficiency= %s,description=%s  WHERE idresourceCyclotron = %s"
        pk = selected_rec[4]
        table_name = 'resourcecyclotron'
        labels = (('Version', ''), ('Capacity', '(mci/h)'), ('Constant Efficiency', '(mCi/mA)'), ('Description', ''))
        save_title = "Save Changes"

        editCyclPopup.edit_popup(labels, selected_rec, save_title, query, pk, cyclo_tabel,table_name)


def deleteCyclotronfun():
    query = "DELETE FROM resourcecyclotron WHERE idresourceCyclotron = %s"
    table_name='resourcecyclotron'
    cyclo_tabel.delete_record(query,table_name)

def addCyclotronfun():
    addCyclPopup = Popup()
    popup_size = "900x550"
    addCyclPopup.open_pop('Add Cyclotron Details',popup_size)
    labels = (('Version', ''), ('Capacity', '(mci/h)'), ('Constant Efficiency', '(mCi/mA)'), ('Description', ''))
    save_title = "Add Cyclotron"
    insertquery = "INSERT INTO resourcecyclotron SET version = %s ,capacity= %s, constant_efficiency= %s,description=%s"
    # selectIDquery = "SELECT MAX(idresourceCyclotron) FROM resourcecyclotron"
    table_name='resourcecyclotron'
    addCyclPopup.add_popup(labels, save_title, insertquery, cyclo_tabel,table_name)

#cyclotron buttons

#Create a button in the main Window to add record - cyclotron
cyclotronAddIcon = Image.open("addIcon.png")
resizedCycloAddIcon = cyclotronAddIcon.resize((25, 25), Image.ANTIALIAS)
imgAddCyclotron = ImageTk.PhotoImage(resizedCycloAddIcon)
addCyclotronButton = Button(cycloSettingsFrame, image=imgAddCyclotron, borderwidth=0, command=lambda : addCyclotronfun())
addCyclotronButton.pack(side= LEFT)
addCyclotronButton.place(x=table_place_x + cyclo_tabel.winfo_reqwidth() -100, y=table_place_y+14)

#Create a button in the main Window to edit  record (open the popup) - cyclotron
cyclotronEditIcon = Image.open("editIcon.jpg")
resizedCycloEditIcon = cyclotronEditIcon.resize((20, 20), Image.ANTIALIAS)
imgEditCyclotron = ImageTk.PhotoImage(resizedCycloEditIcon)
# editCyclotronButton = Button(ctcloSettingsFrame, image=imgEditCyclotron, borderwidth=0, command= lambda :editCyclotronfun())
editCyclotronButton = Button(cycloSettingsFrame, image=imgEditCyclotron, borderwidth=0, command= lambda :editCyclotronfun())

editCyclotronButton.pack(side= LEFT)
editCyclotronButton.place(x=table_place_x + cyclo_tabel.winfo_reqwidth() -50, y=table_place_y+15)


# Create a button in the main Window to Delete record - cyclotron
cyclotronDeleteIcon = Image.open("‏‏deleteIcon.png")
resizedCycloDeleteIcon = cyclotronDeleteIcon.resize((20, 20), Image.ANTIALIAS)
imgDeleteCyclotron = ImageTk.PhotoImage(resizedCycloDeleteIcon)
deleteCyclotronButton = Button(cycloSettingsFrame, image=imgDeleteCyclotron, borderwidth=0, command=lambda : deleteCyclotronfun())
deleteCyclotronButton.pack(side=LEFT)
deleteCyclotronButton.place(x=table_place_x + cyclo_tabel.winfo_reqwidth(), y=table_place_y + 15)


##################### settings new - module #####################
#module frame
moduleSettingsFrame = Frame(root)
# h = Scrollbar(moduleSettingsFrame, orient='horizontal')
# moduleSettingsFrame.pack(fill=X)

# feed label - module
feedLabel = Label(moduleSettingsFrame, text = 'Settings ➝ ', font=label_font_flag,fg=label_color)
PlaceLable_X=50
PlaceLable_Y=10
feedLabel.pack(side=LEFT)
feedLabel.place(x=PlaceLable_X,y=PlaceLable_Y)

feedLabeflag = Label(moduleSettingsFrame, text = 'Module', font=label_font_flag_on_page,fg=label_color)

PlaceLable2_X=135
feedLabeflag.pack(side=LEFT)
feedLabeflag.place(x=PlaceLable2_X,y=PlaceLable_Y)

##################### Module #####################

# Module Details label
moduleLabel = Label(moduleSettingsFrame, text = 'Module Details', font=sub_label_font,fg=label_color)
# module_Lable_place_x=80
# module_Lable_place_y=60

moduleLabel.pack(side=LEFT)
moduleLabel.place(x=Lable_place_x,y=Lable_place_y)
moduleLabel.pack(side=RIGHT)
moduleLabel.place(x=Lable_place_x,y=Lable_place_y)

###module tabel###
scroll_width=20
tab_side=LEFT
x=420
y= 150
frame=moduleSettingsFrame
list_height=5
# table_place_x = 80
# table_place_y=80

columns_name_list=('Version', 'Capacity (mci/h)', 'Description')

queryModule = "SELECT * FROM resourcemodule WHERE ISNULL(deleted)"

module_tabel=table(frame,scroll_width,list_height,tab_side,x,y,table_place_x,
                   table_place_y)
module_tabel.create_fully_tabel( columns_name_list, queryModule)

###module functions###
def editModulefun():
    selected_rec = module_tabel.selected()
    selected_non = module_tabel.selected_is_non(selected_rec)
    if not selected_non:
        # popup_size = "800x450"
        popup_size = "900x550"
        editModulePopup = Popup()
        editModulePopup.open_pop('Edit Module Details', popup_size)

        query = "UPDATE resourcemodule SET version = %s ,capacity= %s, description=%s  WHERE idresourcemodule = %s"
        table_name = 'resourcemodule'
        pk = selected_rec[3]

        labels = (('Version', ''), ('Capacity', '(mci/h)'),  ('Description', ''))
        save_title = "Save Changes"

        editModulePopup.edit_popup(labels, selected_rec, save_title, query, pk, module_tabel,  table_name)


def addModulefun():
    addModulePopup = Popup()
    # popup_size = "800x450"
    popup_size = "900x550"
    addModulePopup.open_pop('Add Module Details',popup_size)
    labels = (('Version', ''), ('Capacity', '(mci/h)'), ('Description', ''))
    save_title = "Add Module"
    insetQuery = "INSERT INTO resourcemodule SET version = %s ,capacity= %s,description=%s"
    table_name='resourcemodule'
    addModulePopup.add_popup(labels, save_title, insetQuery, module_tabel, table_name)

def deleteModulefun():
    query = "DELETE FROM resourcemodule WHERE idresourcemodule = %s"
    table_name='resourcemodule'
    module_tabel.delete_record(query,table_name)


#module buttons

#Create a button in the main Window to add record - module
moduleAddIcon = Image.open("addIcon.png")
resizedModuleAddIcon = moduleAddIcon.resize((25, 25), Image.ANTIALIAS)
imgAddModule = ImageTk.PhotoImage(resizedModuleAddIcon)
addModuleButton = Button(moduleSettingsFrame, image=imgAddModule, borderwidth=0, command=addModulefun)
addModuleButton.pack(side= LEFT)
addModuleButton.place(x=table_place_x+ module_tabel.winfo_reqwidth() -100 , y=table_place_y+14)

#Create a button in the main Window to edit  record (open the popup) - module
moduleEditIcon = Image.open("editIcon.jpg")
resizedModuleEditIcon = moduleEditIcon.resize((20, 20), Image.ANTIALIAS)
imgEditModule = ImageTk.PhotoImage(resizedModuleEditIcon)
editModuleButton = Button(moduleSettingsFrame, image=imgEditModule, borderwidth=0, command=editModulefun)
editModuleButton.pack(side= LEFT)
editModuleButton.place(x=table_place_x+module_tabel.winfo_reqwidth() - 50, y=table_place_y+15)


#Create a button in the main Window to Delete record - module
moduleDeleteIcon = Image.open("‏‏deleteIcon.png")
resizedModuleDeleteIcon = moduleDeleteIcon.resize((20, 20), Image.ANTIALIAS)
imgDeleteModule = ImageTk.PhotoImage(resizedModuleDeleteIcon)
deleteModuleButton = Button(moduleSettingsFrame, image=imgDeleteModule, borderwidth=0, command=deleteModulefun)
deleteModuleButton.pack(side= LEFT)
deleteModuleButton.place(x=table_place_x+module_tabel.winfo_reqwidth(), y=table_place_y+15)



# ##################### Material #####################
##################### settings new - Material #####################
#material frame
materialSettingsFrame = Frame(root)
# h = Scrollbar(materialSettingsFrame, orient='horizontal')
# materialSettingsFrame.pack(fill=X)

# feed label - material
feedLabelmaterial = Label(materialSettingsFrame, text = 'Settings ➝ ', font=label_font_flag,fg=label_color)
PlaceLable_X=50
PlaceLable_Y=10
feedLabelmaterial.pack(side=LEFT)
feedLabelmaterial.place(x=PlaceLable_X,y=PlaceLable_Y)

feedLabeflag = Label(materialSettingsFrame, text = 'Material', font=label_font_flag_on_page,fg=label_color)

PlaceLable2_X=135
feedLabeflag.pack(side=LEFT)
feedLabeflag.place(x=PlaceLable2_X,y=PlaceLable_Y)

##################### material #####################

# material Details label
materialLabel = Label(materialSettingsFrame, text = 'Material Details', font=sub_label_font,fg=label_color)
# material_Lable_place_x=80
# material_Lable_place_y=60

materialLabel.pack(side=LEFT)
materialLabel.place(x=Lable_place_x,y=Lable_place_y)

###material tabel###
scroll_width=20
tab_side=LEFT
x=420
y= 150
frame=materialSettingsFrame
list_height=5
# table_place_x = 80
# table_place_y=80

columns_name_list=[' Material ']

queryMaterial = "SELECT * FROM material WHERE ISNULL(deleted)"

material_tabel=table(frame,scroll_width,list_height,tab_side,x,y,table_place_x,
                     table_place_y)
material_tabel.create_fully_tabel( columns_name_list, queryMaterial)

###material functions###
def editMaterialfun():
    selected_rec = material_tabel.selected()
    selected_non = material_tabel.selected_is_non(selected_rec)
    if not selected_non:
        # popup_size = "800x450"
        popup_size = "700x550"
        editMaterialPopup = Popup()
        editMaterialPopup.open_pop('Edit Material Details', popup_size)

        query = "UPDATE material SET materialName = %s   WHERE idmaterial = %s"
        table_name = 'material'
        pk = selected_rec[1]
        labels = [('Material', '')]
        save_title = "Save Changes"

        editMaterialPopup.edit_popup(labels, selected_rec, save_title, query, pk, material_tabel,  table_name)


def addMaterialfun():
    addMaterialPopup = Popup()
    # popup_size = "800x450"
    popup_size = "900x550"
    addMaterialPopup.open_pop('Add Material Details', popup_size)
    labels = [('Material', '')]
    save_title = "Add Material"
    insetQuery = "INSERT INTO material SET materialName = %s "
    table_name='material'
    addMaterialPopup.add_popup(labels, save_title, insetQuery, material_tabel, table_name)

def deleteMaterialfun():
    query = "DELETE FROM material WHERE idmaterial = %s"
    table_name='material'
    material_tabel.delete_record(query,table_name)


#material buttons

#Create a button in the main Window to add record - material
materialAddIcon = Image.open("addIcon.png")
resizedMaterialAddIcon = materialAddIcon.resize((25, 25), Image.ANTIALIAS)
imgAddMaterial = ImageTk.PhotoImage(resizedMaterialAddIcon)
addMaterialButton = Button(materialSettingsFrame, image=imgAddModule, borderwidth=0, command=addMaterialfun)
addMaterialButton.pack(side= LEFT)
addMaterialButton.place(x=table_place_x + material_tabel.winfo_reqwidth() - 95, y=table_place_y+14)

#Create a button in the main Window to edit  record (open the popup) - material
materialEditIcon = Image.open("editIcon.jpg")
resizedMaterialEditIcon = materialEditIcon.resize((20, 20), Image.ANTIALIAS)
imgEditMaterial = ImageTk.PhotoImage(resizedMaterialEditIcon)
editMaterialButton = Button(materialSettingsFrame, image=imgEditMaterial, borderwidth=0, command=editMaterialfun)
editMaterialButton.pack(side= LEFT)
editMaterialButton.place(x=table_place_x + material_tabel.winfo_reqwidth() - 45, y=table_place_y+15)

#Create a button in the main Window to Delete record - material
materialDeleteIcon = Image.open("‏‏deleteIcon.png")
resizedMaterialDeleteIcon = materialDeleteIcon.resize((20, 20), Image.ANTIALIAS)
imgDeleteMaterial = ImageTk.PhotoImage(resizedMaterialDeleteIcon)
deleteMaterialButton = Button(materialSettingsFrame, image=imgDeleteMaterial, borderwidth=0, command=deleteMaterialfun)
deleteMaterialButton.pack(side= LEFT)
deleteMaterialButton.place(x=table_place_x + material_tabel.winfo_reqwidth() , y=table_place_y+15)


##################### settings new - Hospitals #####################
#hospital frame
hospitalFrame = Frame(root)
# hospitalFrame.pack(fill=X)

# feed label - hospital
feedLabel = Label(hospitalFrame, text = 'Settings ➝ ', font=label_font_flag,fg=label_color)
PlaceLable_X=50
PlaceLable_Y=10
feedLabel.pack(side=LEFT)
feedLabel.place(x=PlaceLable_X,y=PlaceLable_Y)

feedLabeflag = Label(hospitalFrame, text = 'hospital', font=label_font_flag_on_page,fg=label_color)

PlaceLable2_X=135
feedLabeflag.pack(side=LEFT)
feedLabeflag.place(x=PlaceLable2_X,y=PlaceLable_Y)


# hospital Details label
hospitalLabel = Label(hospitalFrame, text = 'Hospitals Details', font=sub_label_font,fg=label_color)
# module_Lable_place_x=80
# module_Lable_place_y=60

hospitalLabel.pack(side=LEFT)
hospitalLabel.place(x=Lable_place_x,y=Lable_place_y)


#hospital table
scroll_width=20
tab_side=LEFT
x=650
y= 160
frame=hospitalFrame
list_height=30
c = 80

lable_place_x = 80
lable_place_y=70

columns_name_list=('        Name        ', 'Fixed Activity Level (mci)', 'Transport Time - min (minutes)', 'Transport Time - man (minutes)')

hospital_query="SELECT * FROM hospital WHERE ISNULL(deleted)"

hospital_tabel=table(frame,scroll_width,list_height,tab_side,x,y,lable_place_x,
                     lable_place_y)
hospital_tabel.create_fully_tabel( columns_name_list, hospital_query)

hospitalFrame.pack(fill='both',expand=1)

###hospital functions###
def editHospitalfun():
    selected_rec = hospital_tabel.selected()
    selected_non = hospital_tabel.selected_is_non(selected_rec)
    if not selected_non:
        editHospitalPopup = Popup()
        # popup_size = "800x450"
        popup_size = "900x550"
        editHospitalPopup.open_pop('Edit Hospital Details',popup_size)
        table_name= 'hospital'
        query = "UPDATE hospital SET Name = %s ,Fixed_activity_level= %s, Transport_time_min=%s ,Transport_time_max=%s WHERE idhospital = %s"

        pk = selected_rec[4]

        labels = (('Name', ''), ('Fixed activity level', '(mci/h)'),  ('Transport time - min', '(min)'),  ('Transport time - max', '(min)'))
        save_title = "Save Changes"

        editHospitalPopup.edit_popup(labels, selected_rec, save_title, query, pk, hospital_tabel,table_name)


def addHospitalfun():
    addHospitalPopup = Popup()
    # popup_size = "800x450"
    popup_size = "900x550"
    addHospitalPopup.open_pop('Add Hospital Details',popup_size)
    labels = (('Name', ''), ('Fixed activity level', '(mci/h)'), ('Transport time - min', '(min)'),('Transport time - max', '(min)'))
    save_title = "Add Hospital"
    insertQuery = "INSERT INTO hospital SET Name = %s ,Fixed_activity_level= %s,Transport_time_min=%s ,Transport_time_max=%s"
    # selectIDquery = "SELECT MAX(idhospital) FROM hospital"
    table_name = 'hospital'
    addHospitalPopup.add_popup(labels, save_title, insertQuery, hospital_tabel, table_name)

def deleteHospitalfun():
    query = "DELETE FROM hospital WHERE idhospital = %s"
    table_name= 'hospital'
    hospital_tabel.delete_record(query,table_name)

#hospital buttons

#Create a button in the main Window to add record - hospital
hospitalAddIcon = Image.open("addIcon.png")
resizedHospitalAddIcon = hospitalAddIcon.resize((25, 25), Image.ANTIALIAS)
imgAddHospital = ImageTk.PhotoImage(resizedHospitalAddIcon)
addHospitalButton = Button(hospitalFrame, image=imgAddHospital, borderwidth=0, command=lambda : addHospitalfun())
addHospitalButton.pack(side= LEFT)
addHospitalButton.place(x=lable_place_x + hospital_tabel.winfo_reqwidth() - 100, y=lable_place_y+14)


#Create a button in the main Window to edit  record (open the popup) - hospital
hospitalEditIcon = Image.open("editIcon.jpg")
resizedHospitalEditIcon = hospitalEditIcon.resize((20, 20), Image.ANTIALIAS)
imgEditHospital = ImageTk.PhotoImage(resizedHospitalEditIcon)
editHospitalButton = Button(hospitalFrame, image=imgEditHospital, borderwidth=0, command= lambda :editHospitalfun())

editHospitalButton.pack(side= LEFT)
editHospitalButton.place(x=lable_place_x + hospital_tabel.winfo_reqwidth() - 50, y=lable_place_y+15)


# Create a button in the main Window to Delete record - hospital
hospitalDeleteIcon = Image.open("‏‏deleteIcon.png")
resizedHospitalDeleteIcon = hospitalDeleteIcon.resize((20, 20), Image.ANTIALIAS)
imgDeleteHospital = ImageTk.PhotoImage(resizedHospitalDeleteIcon)
deleteHospitalButton = Button(hospitalFrame, image=imgDeleteHospital, borderwidth=0, command=lambda : deleteHospitalfun())
deleteHospitalButton.pack(side=LEFT)

deleteHospitalButton.place(x=lable_place_x + hospital_tabel.winfo_reqwidth() , y=lable_place_y + 15)



################work plan page###############################################################

#################### Work Plan Page #####################
#Work Plan frame
WorkPlanFrame = Frame(root)
# h = Scrollbar(WorkPlanFrame, orient='horizontal')
WorkPlanFrame.pack(fill=X)

##################### Work Plan #####################
# Work Plan Details label
WorkPlanLabel = Label(WorkPlanFrame, text = 'Work Plans', font='Helvetica',fg='#034672')
Lable_place_x=80
Lable_place_y=60

WorkPlanLabel.pack(side=LEFT)
WorkPlanLabel.place(x=Lable_place_x,y=Lable_place_y)

###Work Plan tabel###
scroll_width=20
tab_side=LEFT
x=330
y= 140
frame=WorkPlanFrame
list_height=5
table_place_x = 80
table_place_y = 80
columns_name_list=('    Date   ',' Material ' )
query = "SELECT WP.idworkplan, WP.Date, m.materialName FROM workplan WP JOIN material M ON WP.materialID=M.idmaterial"

wp_tabel=table(frame,scroll_width,list_height,tab_side,x,y,table_place_x,
               table_place_y,)
wp_tabel.create_fully_tabel( columns_name_list, query)


# ###Work Plan functions###
# def editWPfun():
#     selected_rec = cyclo_tabel.selected()
#     selected_non=cyclo_tabel.selected_is_non(selected_rec)
#     if not selected_non:
#         editCyclPopup = Popup()
#         editCyclPopup.open_pop('Edit Cyclotron Details')
#
#         query = "UPDATE resourcecyclotron SET version = %s ,capacity= %s, constant_efficiency= %s,description=%s  WHERE idresourceCyclotron = %s"
#         pk = selected_rec[4]
#         table_name = 'resourcecyclotron'
#         labels = (('Version', ''), ('Capacity', '(mci/h)'), ('Constant Efficiency', '(mCi/mA)'), ('Description', ''))
#         save_title = "Save Changes"
#
#         editCyclPopup.edit_popup(labels, selected_rec, save_title, query, pk, cyclo_tabel,table_name)
#
#
# def deleteCyclotronfun():
#     query = "DELETE FROM resourcecyclotron WHERE idresourceCyclotron = %s"
#     table_name='resourcecyclotron'
#     cyclo_tabel.delete_record(query,table_name)
#
def addWPfun():
    addWPPopup = Popup()
    popup_size = "800x450"
    addWPPopup.open_pop('Create Work Plan',popup_size)
    addWPPopup.add_wp_popup()

# #work plan buttons
# #Create a button in the main Window to edit  record (open the popup) - work plan
# wpEditIcon = Image.open("editIcon.jpg")
# resizedWPEditIcon = wpEditIcon.resize((20, 20), Image.ANTIALIAS)
# imgEditwp = ImageTk.PhotoImage(resizedWPEditIcon)
# # editCyclotronButton = Button(ctcloSettingsFrame, image=imgEditCyclotron, borderwidth=0, command= lambda :editCyclotronfun())
# editWPButton = Button(WorkPlanFrame, image=imgEditwp, borderwidth=0, command= lambda :editCyclotronfun())
#
# editWPButton.pack(side= LEFT)
# editWPButton.place(x=table_place_x+450, y=table_place_y+15)
#
#Create a button in the main Window to add record - work plan
wpAddIcon = Image.open("addIcon.png")
resizedWPAddIcon = wpAddIcon.resize((25, 25), Image.ANTIALIAS)

imgAddWP = ImageTk.PhotoImage(resizedWPAddIcon)
addWPButton = Button(WorkPlanFrame, image=imgAddWP, borderwidth=0, command=lambda : addWPfun())
addWPButton.pack(side= LEFT)
addWPButton.place(x=table_place_x+160, y=table_place_y+14)
#
#
# # Create a button in the main Window to Delete record - work plan
# wpDeleteIcon = Image.open("‏‏deleteIcon.png")
# resizedWPDeleteIcon = wpDeleteIcon.resize((20, 20), Image.ANTIALIAS)
# imgDeleteWP = ImageTk.PhotoImage(resizedWPDeleteIcon)
# deleteWPButton = Button(WorkPlanFrame, image=imgDeleteWP, borderwidth=0, command=lambda : deleteCyclotronfun())
# deleteWPButton.pack(side=LEFT)
# deleteWPButton.place(x=table_place_x + 500, y=table_place_y + 15)

################### batches #################
#################### batch Page #####################
#batch frame
batchFrame = Frame(root)
# h = Scrollbar(batchFrame, orient='horizontal')
batchFrame.pack(fill=X)

# batch Details label
BatchLabel = Label(batchFrame, text = 'Batches', font=sub_label_font,fg=label_color)
Lable_place_x=80
Lable_place_y=60

BatchLabel.pack(side=LEFT)
BatchLabel.place(x=Lable_place_x,y=Lable_place_y)

#batches table
scroll_width=20
tab_side=LEFT
x=650
y= 160
frame=batchFrame
list_height=30
c = 80

lable_place_x = 80
lable_place_y=70

columns_name_list=('  Date  ', 'Batch Number','Material','TargetCurrentLB ', 'DecayCorrected_TTA (mci)', 'EOS_activity')

batch_query="""SELECT  b.idbatch , wp.Date ,b.batchNumber, m.materialName, b.TargetCurrentLB ,b.DecayCorrected_TTA, b.EOS_activity 
                FROM batch b 
                JOIN workplan wp ON wp.idworkplan = b.workplanID 
                JOIN material m ON m.idmaterial = wp.materialID"""

batch_tabel=table(frame,scroll_width,list_height,tab_side,x,y,lable_place_x,
                  lable_place_y)
batch_tabel.create_fully_tabel( columns_name_list, batch_query)

batchFrame.pack(fill='both',expand=1)

###batch functions###
def editBatchfun():
    selected_rec = batch_tabel.selected()
    selected_non = batch_tabel.selected_is_non(selected_rec)
    if not selected_non:
        editBatchPopup = Popup()
        # popup_size = "800x450"
        popup_size = "900x550"
        editBatchPopup.open_pop('Edit Batch Details',popup_size)
        table_name= 'batch'
        query = "UPDATE batch SET TargetCurrentLB = %s ,DecayCorrected_TTA= %s, EOS_activity=%s  WHERE idbatch = %s"

        pk = selected_rec[6]

        labels = (('TargetCurrentLB', ''), ('DecayCorrected_TTA', '(mci/h)'),  ('EOS_activity', ''))
        save_title = "Save Changes"

        # def edit_popup(self, labels, valueList, save_title, *args, table_name):
        editBatchPopup.edit_popup(labels, selected_rec, save_title, query, pk, batch_tabel,table_name)

#batch buttons

#Create a button in the main Window to edit  record (open the popup) - hospital
batchEditIcon = Image.open("editIcon.jpg")
resizedBatchEditIcon = batchEditIcon.resize((20, 20), Image.ANTIALIAS)
imgEditBatch = ImageTk.PhotoImage(resizedBatchEditIcon)
editBatchButton = Button(batchFrame, image=imgEditBatch, borderwidth=0, command= lambda :editBatchfun())

editBatchButton.pack(side= LEFT)
editBatchButton.place(x=lable_place_x + batch_tabel.winfo_reqwidth() - 50, y=lable_place_y+15)

# workPlanButton.config(bg='#F0F0F0');  ##F0F0F0 is default color(gray)
# ordersButton.config(bg='#F0F0F0');
ordersFrame.forget();
cycloSettingsFrame.forget()
moduleSettingsFrame.forget()
hospitalFrame.forget()
WorkPlanFrame.forget()
batchFrame.forget()
################################and of work plan page###############################


#####################New toolbar #####################

toolbarbgcolor = "white"
toolbar = Frame(root, bg=toolbarbgcolor)
toolbar.grid(sticky='nesw')

# logo - toolbar
LogoImagePath = Image.open("LogoImage.png")
LogoImageResize = LogoImagePath.resize((120, 57),Image.ANTIALIAS)
LogoImage = ImageTk.PhotoImage(LogoImageResize)
Label(toolbar,image=LogoImage).pack(side=LEFT,padx=10,pady=6)


#toolbar function
def work_plan_page():
    ordersButton.config(bg='#F0F0F0');# Orders uttun color default value
    mbtn.config(bg='#F0F0F0'); ##F0F0F0 is default color(gray)
    workPlanButton.config(bg="gray");
    WorkPlanFrame.pack(fill=X);
    ordersFrame.forget();
    moduleSettingsFrame.forget();
    materialSettingsFrame.forget();
    hospitalFrame.forget();
    cycloSettingsFrame.forget();

# work plan button - toolbar
workPlanButton = Button(toolbar, text="Work Plans",font='Helvetica 11',  command=lambda: work_plan_page())
workPlanButton.pack(side=LEFT,padx=10,pady=3)


# Orders button - toolbar
ordersButton = Button (toolbar, text="Orders", font='Helvetica 11',command=Orders_page)
ordersButton.pack(side=LEFT,padx=10,pady=3)

# # Batches button - toolbar
# ordersButton = Button (toolbar, text="Batches", font='Helvetica 11')
# ordersButton.pack(side=LEFT,padx=10,pady=3)


# Reports button - toolbar
reportsButton = Button (toolbar, text="Reports", font='Helvetica 11')
reportsButton.pack(side=LEFT,padx=10,pady=3)

# settings Icon - toolbar
global imgSettings;
settingsIcon = Image.open("D:\PythonProjects\Cyclotron\gearIcon.png")
resizedSettingsIcon = settingsIcon.resize((35,35), Image.ANTIALIAS)
imgSettings = ImageTk.PhotoImage(resizedSettingsIcon,master=toolbar)
# Button(toolbar, image=imgSettings, borderwidth=0).pack(side=RIGHT,padx=10,pady=3)
mbtn = Menubutton(toolbar, image=imgSettings, borderwidth=0)
mbtn.pack(side=RIGHT,padx=10,pady=3)
mbtn.menu = Menu(mbtn, tearoff = 0)
mbtn["menu"] = mbtn.menu
selected_settings_option = StringVar()

def menu_item_selected(label):
    mbtn.config(bg="gray");
    ordersButton.config(bg='#F0F0F0');#Default value
    workPlanButton.config(bg='#F0F0F0');#Default value
    if label == 'Cyclotron':

        cycloSettingsFrame.pack(fill=X);
        ordersFrame.forget();
        moduleSettingsFrame.forget()
        materialSettingsFrame.forget()
        hospitalFrame.forget()
        WorkPlanFrame.forget()


    elif label == 'Module':
        moduleSettingsFrame.pack(fill=X);
        ordersFrame.forget();
        cycloSettingsFrame.forget();
        materialSettingsFrame.forget();
        hospitalFrame.forget();
        WorkPlanFrame.forget();

    elif label == 'Hospital':
        hospitalFrame.pack(fill=X);
        ordersFrame.forget();
        cycloSettingsFrame.forget();
        materialSettingsFrame.forget();
        moduleSettingsFrame.forget();
        WorkPlanFrame.forget();

    else:
        materialSettingsFrame.pack(fill=X);
        ordersFrame.forget();
        cycloSettingsFrame.forget()
        moduleSettingsFrame.forget()
        hospitalFrame.forget()
        WorkPlanFrame.forget()


selected_settings_option.trace("w", menu_item_selected)

mbtn.menu.add_radiobutton(label="Cyclotron", command= lambda: menu_item_selected("Cyclotron"))
mbtn.menu.add_radiobutton(label="Module", command= lambda: menu_item_selected("Module"))
mbtn.menu.add_radiobutton(label="Material", command= lambda: menu_item_selected("Material"))
mbtn.menu.add_radiobutton(label="Hospital", command= lambda: menu_item_selected("Hospital"))



# print(mbtn.selection_get())
toolbar.pack(side=TOP, fill=X)

toolbar.grid_columnconfigure(1, weight=1)


#########################################################################
# def setting_page():
#     """ this function is swap function for viewing setting frame/page"""
#     #hospitalsButton.config(bg='#F0F0F0')
#     #ordersButton.config(bg='#F0F0F0');
#     #settingButton.config(bg="gray");
#     #SettingsFrame.pack(fill='both',expand=1);
#     hospitalFrame.forget();
#     ordersFrame.forget();

# settings Icon - toolbar

# settingsIcon = Image.open("D:\PythonProjects\Cyclotron\gearIcon.png")
# resizedSettingsIcon = settingsIcon.resize((35, 35), Image.ANTIALIAS)
# imgSettings = ImageTk.PhotoImage(resizedSettingsIcon,master=toolbar)
# settingButton=Button(toolbar,command=setting_page ,image=imgSettings,activebackground='red', borderwidth=0)
# settingButton.pack(side=RIGHT, padx=10, pady=3)


toolbar.pack(side=TOP, fill=X)

toolbar.grid_columnconfigure(1, weight=1)

#####################################Old Settings frame####################
#
#
# SettingsFrame = Frame(root)
# #h = Scrollbar(SettingsFrame, orient='horizontal')
# SettingsFrame.pack(fill='both',expand=1)
#
# # feed label
# feedLabel = Label(SettingsFrame, text='Settings', font=('Helvetica', 26, 'bold'), fg='#034672')
# PlaceLable_X = 50
# PlaceLable_Y = 10
#
# feedLabel.pack(side=LEFT)
# feedLabel.place(x=PlaceLable_X, y=PlaceLable_Y)
#
# # admin label
# AdminLabelSettingpage1 = Label(SettingsFrame, text=f"{Permission.ValidateTypeOfUser} connected:", font=('Helvetica', 13, 'bold'), fg='red')
# AdminLabelSettingpage1.pack();
# AdminLabelSettingpage1.place(x=350, y=20);
#
# # admin connected label
# NameOfAdminLabeLSettingspage2 = Label(SettingsFrame, text=Permission.user_verified, font=('Helvetica', 13, 'bold'), fg='red')
# NameOfAdminLabeLSettingspage2.pack();
# NameOfAdminLabeLSettingspage2.place(x=510, y=20);
#
#
# ##################### Cyclotron #####################
# # Cyclotron Details label
# CyclotronLabel = Label(SettingsFrame, text='Cyclotron Details', font=('Helvetica', 15, 'bold'), fg='#034672')
# cyclo_Lable_place_x = 80
# cyclo_Lable_place_y = 70
#
# CyclotronLabel.pack(side=LEFT)
# CyclotronLabel.place(x=cyclo_Lable_place_x, y=cyclo_Lable_place_y)
#
# # scrollbar
# Cyclotron_scroll = Scrollbar(SettingsFrame, orient="vertical", width=20)
# Cyclotron_scroll.pack(side=LEFT)
# Cyclotron_scroll.place(x=613, y=160)
#
# cyclo_list = ttk.Treeview(SettingsFrame, yscrollcommand=Cyclotron_scroll.set, height=5)
#
# cyclo_list.pack(side=LEFT, padx=cyclo_Lable_place_x + 30, pady=cyclo_Lable_place_y + 50)
#
# # Cyclotron_scroll.config(command=cyclo_list.yview)
# # Cyclotron_scroll.config(command=cyclo_list.xview)
#
# # column define
#
# cyclo_list['columns'] = ('Version', 'Capacity (mci/h)', 'Constant Efficiency (mCi/mA)', 'Description')
#
# # column format
# width_Version = 90
# width_Capacity = 110
# width_Efficiency = 185
# width_Description = 110
#
# cyclo_list.column("#0", width=0, stretch=NO)
# cyclo_list.column("Version", anchor=CENTER, width=width_Version)
# cyclo_list.column("Capacity (mci/h)", anchor=CENTER, width=width_Capacity)
# cyclo_list.column("Constant Efficiency (mCi/mA)", anchor=CENTER, width=width_Efficiency)
# cyclo_list.column("Description", anchor=CENTER, width=width_Description)
#
# # Create Headings
# cyclo_list.heading("#0", text="", anchor=CENTER)
# cyclo_list.heading("Version", text="Version", anchor=CENTER)
# cyclo_list.heading("Capacity (mci/h)", text="Capacity (mci/h)", anchor=CENTER)
# cyclo_list.heading("Constant Efficiency (mCi/mA)", text="Constant Efficiency (mCi/mA)", anchor=CENTER)
# cyclo_list.heading("Description", text="Description", anchor=CENTER)
#
# # add data from db
# cursor = db.cursor()
# # Test to see if DB was created
# # cursor.execute("SHOW DATABASES")
# # for cyclotron in cursor:
# #     print(cyclotron)
#
# cursor.execute("SELECT * FROM resourcecyclotron");
# cyclotrons = cursor.fetchall();
#
# #Insert data of Settings page into My-SQl
# #The INSERT IGNORE statement will cause MySQL to do nothing when the insertion throws an error. If there’s no error, then a new row will be added to the table.
# cursor.execute("INSERT IGNORE INTO resourcecyclotron (idresourceCyclotron,version,capacity,constant_efficiency,description) VALUES (1,2.1,2000,220,'site1'),(2,2.2,1700,150,'site2');")
# #cleanup
# db.commit()
# # cursor.close()
# # db.close()
#
#
# iid = 0
# for cyclo in cyclotrons:
#     print(cyclo)
#     cyclo_list.insert(parent='', index='end', iid=iid, text='',
#                       values=(cyclo[1], cyclo[2], cyclo[3], cyclo[4]))
#     iid += 1
#
# cyclo_list.pack()
#
# # frame = Frame(root)
# # frame.pack()
#
# get_version = ""
# get_capacity = ""
# get_efficiency = ""
# get_description = ""
#
#
# class Cyclotron:
#     def _init_(self, version, capacity, constant_efficiency, description):
#         self.version = version
#         self.capacity = capacity
#         self.constant_efficiency = constant_efficiency
#         self.description = description
#
#     # def edit(self):
#     #    def update_record(get_version, get_capacity, get_efficiency, get_description):
#     #       selected = cyclo_list.focus()
#     #       # save new data
#     #       print("get_version" + get_version)
#     #       cyclo_list.item(selected, text="", values=(get_version, get_capacity, get_efficiency, get_description))
#     #
#     #       # # clear entry boxes
#     #       # Version_entry.delete(0, END)
#     #       # Capacity_entry.delete(0, END)
#     #       # Efficiency_entry.delete(0, END)
#
#
# def open_popup_cyclotron():
#     edit_popup = Toplevel(root)
#     edit_popup.geometry("900x400")
#     edit_popup.title("Edit Cyclotron Details")
#     Label(edit_popup, text="Edit Cyclotron Details", font=('Helvetica 17 bold'), fg='#034672').place(x=10, y=18)
#
#     # labels
#     popup_label_y = 80
#     Version = Label(edit_popup, text="Version")
#     Version.grid(row=1, column=1)
#     version_x = 20
#     Version.place(x=version_x, y=popup_label_y)
#
#     Capacity = Label(edit_popup, text="Capacity")
#     Capacity_units = Label(edit_popup, text="(mci/h)")
#     Capacity_units.config(font=("Courier", 9))
#     Capacity.grid(row=1, column=2)
#     capacity_x = version_x + Version.winfo_reqwidth() + 70
#     Capacity.place(x=capacity_x, y=popup_label_y)
#     capacity_units_x = capacity_x + Capacity.winfo_reqwidth()
#     Capacity_units.place(x=capacity_units_x, y=popup_label_y + 7)
#
#     Efficiency = Label(edit_popup, text="Constant Efficiency")
#     Efficiency_units = Label(edit_popup, text="(mCi/mA)")
#     Efficiency_units.config(font=("Courier", 9))
#     Efficiency.grid(row=1, column=3)
#     efficiency_x = capacity_units_x + Capacity_units.winfo_reqwidth() + 50
#     Efficiency.place(x=efficiency_x, y=popup_label_y)
#     efficiency_units_x = efficiency_x + Efficiency.winfo_reqwidth()
#     Efficiency_units.place(x=efficiency_units_x, y=popup_label_y + 7)
#
#     Description = Label(edit_popup, text="Description")
#     Description.grid(row=1, column=3)
#     description_x = efficiency_units_x + Efficiency_units.winfo_reqwidth() + 30
#     Description.place(x=description_x, y=popup_label_y)
#
#     # Entry boxes
#     Version_entry = Entry(edit_popup, width=10)
#     Version_entry.grid(row=2, column=1)
#     Version_entry.place(x=version_x + 3, y=popup_label_y + 30)
#
#     Capacity_entry = Entry(edit_popup, width=14)
#     Capacity_entry.grid(row=2, column=2)
#     Capacity_entry.place(x=capacity_x, y=popup_label_y + 30)
#
#     Efficiency_entry = Entry(edit_popup, width=15)
#     Efficiency_entry.grid(row=2, column=3)
#     Efficiency_entry.place(x=efficiency_x, y=popup_label_y + 30)
#
#     Description_entry = Entry(edit_popup, width=15)
#     Description_entry.grid(row=2, column=4)
#     Description_entry.place(x=description_x, y=popup_label_y + 30)
#
#     # # clear entry boxes
#     # Version_entry.delete(0, END)
#     # Capacity_entry.delete(0, END)
#     # Efficiency_entry.delete(0, END)
#
#     # grab record
#     selected = cyclo_list.focus()
#     # grab record values
#     values = cyclo_list.item(selected, 'values')
#     # temp_label.config(text=selected)
#
#     # insert cyclotron details from db to entry boxes
#     Version_entry.insert(0, values[0])
#     Capacity_entry.insert(0, values[1])
#     Efficiency_entry.insert(0, values[2])
#     Description_entry.insert(0, values[3])
#
#     # get_version = Version_entry.get()
#     # get_capacity = Capacity_entry.get()
#     # get_efficiency = Efficiency_entry.get()
#     # get_description = Description_entry.get()
#
#     select_button = Button(edit_popup, text="Save Changes",
#                            command=lambda: update_record(Version_entry.get(), Capacity_entry.get(),
#                                                          Efficiency_entry.get(), Description_entry.get()))
#     select_button.pack(side=LEFT)
#     select_button.place(x=370, y=250)
#
#
# # in the class?
# def update_record(get_version, get_capacity, get_efficiency, get_description):
#     print("get_version" + get_version)
#     selected = cyclo_list.focus()
#     print(cyclo_list.item(selected, 'values'))
#     # save new data
#     cyclo_list.item(selected, text="", values=(get_version, get_capacity, get_efficiency, get_description))
#
#     # # clear entry boxes
#     # Version_entry.delete(0, END)
#     # Capacity_entry.delete(0, END)
#     # Efficiency_entry.delete(0, END)
#
#
# # # save Record
# # def update_record():
# #     selected = cyclo_list.focus()
# #     # save new data
# #     cyclo_list.item(selected, text="", values=(Version_entry.get(), Capacity_entry.get(), Efficiency_entry.get()))
# #
# #     # clear entry boxes
# #     Version_entry.delete(0, END)
# #     Capacity_entry.delete(0, END)
# #     Efficiency_entry.delete(0, END)
#
#
# # Create a button in the main Window to edit  record (open the popup) - cyclotron
# cyclotronEditIcon = Image.open("D:\PythonProjects\Cyclotron\editIcon.jpg")
# resizedCycloEditIcon = cyclotronEditIcon.resize((20, 20), Image.ANTIALIAS)
# imgEditCyclotron = ImageTk.PhotoImage(resizedCycloEditIcon,master=SettingsFrame)
# editCyclotronButton = Button(SettingsFrame, image=imgEditCyclotron, borderwidth=0, command=open_popup_cyclotron)
# editCyclotronButton.pack(side=LEFT)
# editCyclotronButton.place(x=cyclo_Lable_place_x + 450, y=cyclo_Lable_place_y + 15)
#
# # edit_button = Button(SettingsFrame, text= "Edit", command= open_popup_cyclotron, width=4, height=1)
# # edit_button.pack(side= LEFT)
# # edit_button.place(x=270, y=265)
#
#
# # Create a button in the main Window to Delete record - cyclotron
# cyclotronDeleteIcon = Image.open("D:\PythonProjects\Cyclotron\‏‏deleteIcon.png")
# resizedCycloDeleteIcon = cyclotronDeleteIcon.resize((20, 20), Image.ANTIALIAS)
# imgDeleteCyclotron = ImageTk.PhotoImage(resizedCycloDeleteIcon,master=SettingsFrame);
# deleteCyclotronButton = Button(SettingsFrame, image=imgDeleteCyclotron, borderwidth=0, command=open_popup_cyclotron)
# deleteCyclotronButton.pack(side=LEFT)
# deleteCyclotronButton.place(x=cyclo_Lable_place_x + 500, y=cyclo_Lable_place_y + 15)
#
# # Create a button in the main Window to add record - cyclotron
# cyclotronAddIcon = Image.open("D:/PythonProjects/Cyclotron/addIcon.png")
# resizedCycloAddIcon = cyclotronAddIcon.resize((25, 25), Image.ANTIALIAS)
# imgAddCyclotron = ImageTk.PhotoImage(resizedCycloAddIcon,master=SettingsFrame);
# addCyclotronButton = Button(SettingsFrame, image=imgAddCyclotron, borderwidth=0, command=open_popup_cyclotron)
# addCyclotronButton.pack(side=LEFT)
# addCyclotronButton.place(x=cyclo_Lable_place_x + 400, y=cyclo_Lable_place_y + 14)
#
# # add_button = Button(SettingsFrame, text="Add Cyclotron", command= open_popup_cyclotron, width = 4, height=1)
# # add_button.pack(side= LEFT)
# # add_button.place(x=370, y=265)
#
#
#
#
# ##################### Module #####################
# # Module Details label
# moduleLabel = Label(SettingsFrame, text = 'Module Details', font=('Helvetica',15, 'bold'),fg='#034672')
# module_Lable_place_x=700
# module_Lable_place_y=70
#
# moduleLabel.pack(side=RIGHT)
# moduleLabel.place(x=module_Lable_place_x,y=module_Lable_place_y)
#
#
# # scrollbar
# Module_scroll = Scrollbar(SettingsFrame ,orient="vertical",width=20)
# Module_scroll.pack(side=RIGHT)
# Module_scroll.place(x=1035, y= 160)
#
# module_list = ttk.Treeview(SettingsFrame, yscrollcommand=Module_scroll.set,height=5)
#
# module_list.pack(side=LEFT, padx=0, pady=module_Lable_place_y+50)
#
#
#
# # Module_scroll.config(command=cyclo_list.yview)
# # Module_scroll.config(command=cyclo_list.xview)
#
# # column define
#
# module_list['columns'] = ('Version', 'Capacity (mci/h)', 'Description')
#
#
# module_list.column("#0", width=0, stretch=NO)
# module_list.column("Version", anchor=CENTER, width=width_Version)
# module_list.column("Capacity (mci/h)", anchor=CENTER, width=width_Capacity)
# module_list.column("Description", anchor=CENTER, width=width_Description)
#
# # Create Headings
# module_list.heading("#0", text="", anchor=CENTER)
# module_list.heading("Version", text="Version", anchor=CENTER)
# module_list.heading("Capacity (mci/h)", text="Capacity (mci/h)", anchor=CENTER)
# module_list.heading("Description", text="Description", anchor=CENTER)
#
# # add data from db
# cursor = db.cursor()
# cursor.execute("SELECT * FROM resourcemodule")
# modules = cursor.fetchall()
#
# iid=0
# for module in modules:
#     print(module)
#     cyclo_list.insert(parent='', index='end', iid=iid, text='',
#                values=(module[1], module[2], module[3]))
#     iid +=1
#
# module_list.pack()
#
# get_version=""
# get_capacity=""
# get_efficiency=""
# get_description=""
#
# def open_popup_module():
#    edit_popup= Toplevel(root)
#    edit_popup.geometry("900x400")
#    edit_popup.title("Edit Module Details")
#    Label(edit_popup, text= "Edit Module Details", font=('Helvetica 17 bold'), fg='#034672').place(x=10,y=18)
#
#    # labels
#    popup_label_y=80
#    Version = Label(edit_popup, text="Version")
#    Version.grid(row=1, column=1)
#    version_x = 20
#    Version.place(x=version_x, y=popup_label_y)
#
#
#    Capacity = Label(edit_popup, text="Capacity")
#    Capacity_units = Label(edit_popup, text="(mci/h)")
#    Capacity_units.config(font=("Courier", 9))
#    Capacity.grid(row=1, column=2)
#    capacity_x = version_x+Version.winfo_reqwidth()+70
#    Capacity.place(x=capacity_x, y=popup_label_y)
#    capacity_units_x=capacity_x + Capacity.winfo_reqwidth()
#    Capacity_units.place(x=capacity_units_x, y=popup_label_y+7)
#
#
#    Description = Label(edit_popup, text="Description")
#    Description.grid(row=1, column=3)
#    description_x = capacity_units_x + Capacity_units.winfo_reqwidth() + 50
#    Description.place(x=description_x, y=popup_label_y)
#
#    # Entry boxes
#    Version_entry = Entry(edit_popup, width=10)
#    Version_entry.grid(row=2, column=1)
#    Version_entry.place(x=version_x+3, y=popup_label_y+30)
#
#    Capacity_entry = Entry(edit_popup, width=14)
#    Capacity_entry.grid(row=2, column=2)
#    Capacity_entry.place(x=capacity_x, y=popup_label_y+30)
#
#
#    Description_entry = Entry(edit_popup,width=15)
#    Description_entry.grid(row=2, column=4)
#    Description_entry.place(x=description_x, y=popup_label_y+30)
#
#
#    # clear entry boxes
#    Version_entry.delete(0, END)
#    Capacity_entry.delete(0, END)
#    Description_entry.delete(0, END)
#
#    # grab record
#    selected = module_list.focus()
#    # grab record values
#    values = module_list.item(selected, 'values')
#    # temp_label.config(text=selected)
#
#    # output to entry boxes
#    Version_entry.insert(0, values[0])
#    Capacity_entry.insert(0, values[1])
#    Description_entry.insert(0, values[2])
#
#    get_version = Version_entry.get()
#    print(get_version)
#    get_capacity = Capacity_entry.get()
#    get_description = Description_entry.get()
#
#    select_button = Button(edit_popup, text="Save Changes", command=update_record)
#    select_button.pack(side=LEFT)
#    select_button.place(x=370, y=250)
#
#
# #Create a button in the main Window to edit  record (open the popup) - module
# moduleEditIcon = Image.open("D:\PythonProjects\Cyclotron\editIcon.jpg")
# resizedModuleEditIcon = moduleEditIcon.resize((20, 20), Image.ANTIALIAS)
# imgEditModule = ImageTk.PhotoImage(resizedModuleEditIcon,master=SettingsFrame);
# editModuleButton = Button(SettingsFrame, image=imgEditModule, borderwidth=0, command=open_popup_module)
# editModuleButton.pack(side= LEFT)
# editModuleButton.place(x=module_Lable_place_x+250, y=module_Lable_place_y+15)
#
# # edit_button = Button(SettingsFrame, text= "Edit", command= open_popup_module)
# # edit_button.pack(side= LEFT)
# # edit_button.place(x=790, y=270)
#
# #Create a button in the main Window to Delete record - module
# moduleDeleteIcon = Image.open("D:\PythonProjects\Cyclotron\‏‏deleteIcon.png")
# resizedModuleDeleteIcon = moduleDeleteIcon.resize((20, 20), Image.ANTIALIAS)
# imgDeleteModule = ImageTk.PhotoImage(resizedModuleDeleteIcon,master=SettingsFrame);
# deleteModuleButton = Button(SettingsFrame, image=imgDeleteModule, borderwidth=0, command=open_popup_module)
# deleteModuleButton.pack(side= LEFT)
# deleteModuleButton.place(x=module_Lable_place_x+300, y=module_Lable_place_y+15)
#
# #Create a button in the main Window to add record - module
# moduleAddIcon = Image.open("D:/PythonProjects/Cyclotron/addIcon.png")
# resizedModuleAddIcon = moduleAddIcon.resize((25, 25), Image.ANTIALIAS)
# imgAddModule = ImageTk.PhotoImage(resizedModuleAddIcon,master=SettingsFrame);
# addModuleButton = Button(SettingsFrame, image=imgAddModule, borderwidth=0, command=open_popup_cyclotron)
# addModuleButton.pack(side= LEFT)
# addModuleButton.place(x=module_Lable_place_x+200, y=module_Lable_place_y+14)
#
# # add_button = Button(SettingsFrame, text="Add", command= open_popup_cyclotron)
# # add_button.pack(side= LEFT)
# # add_button.place(x=890, y=270)
#
# def update_record():
#     selected = cyclo_list.focus()
#     # save new data
#     print("get_version"+get_version)
#     cyclo_list.item(selected, text="", values=(get_version, get_capacity, get_efficiency, get_description))
#
#     # # clear entry boxes
#     # Version_entry.delete(0, END)
#     # Capacity_entry.delete(0, END)
#     # Efficiency_entry.delete(0, END)
#
#
# # # # save Record
# # # def update_record():
# # #     selected = cyclo_list.focus()
# # #     # save new data
# # #     cyclo_list.item(selected, text="", values=(Version_entry.get(), Capacity_entry.get(), Efficiency_entry.get()))
# # #
# # #     # clear entry boxes
# # #     Version_entry.delete(0, END)
# # #     Capacity_entry.delete(0, END)
# # #     Efficiency_entry.delete(0, END)
#
#
# SettingsFrame.pack(fill='both',expand=1)

root.mainloop()
